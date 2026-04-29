"""
Genera il documento Word della tesi di laurea triennale in Informatica
sull'argomento del progetto BitM-LLM.

Output: tesi_BitM_LLM.docx
"""

import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGDIR = Path(__file__).parent / "tesi_figures"
FIGDIR.mkdir(exist_ok=True)

TEMPLATE_PATH = Path(__file__).parent / "Template_Tesi.docx"

plt.rcParams.update({
    "font.family": "serif",
    "font.size": 10,
    "axes.titlesize": 11,
    "axes.labelsize": 10,
    "axes.spines.top": False,
    "axes.spines.right": False,
    "savefig.dpi": 150,
    "savefig.bbox": "tight",
    "figure.facecolor": "white",
})


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                border.set(qn(f"w:{k}"), str(v))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_break(doc):
    doc.add_page_break()


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
    return h


def add_par(doc, text, bold=False, italic=False, align=None, size=12,
            first_line_indent=True, space_after=6):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_figure(doc, path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))

    # Caption style del template: già configurato (corsivo, ridotto)
    try:
        cap = doc.add_paragraph(style="Caption")
    except KeyError:
        cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(10)


# ── Chart generators ─────────────────────────────────────────────────────────

def fig_pipeline_diagram():
    out = FIGDIR / "pipeline.png"
    fig, ax = plt.subplots(figsize=(11, 2.6))
    stages = [
        ("HTTP\nRequest", "#5B8FF9"),
        ("GeoIP\nMiddleware", "#5AD8A6"),
        ("Rate\nCheck", "#5AD8A6"),
        ("Session\nLoad", "#5AD8A6"),
        ("Feature\nExtractor", "#F6BD16"),
        ("Fast\nRules", "#E8684A"),
        ("LLM\nScorer", "#9270CA"),
        ("Trajectory\nLLM", "#9270CA"),
        ("Policy\nDecide", "#E8684A"),
        ("Response\n+ Log + WS", "#5B8FF9"),
    ]
    n = len(stages)
    box_w, box_h = 1.05, 0.85
    y = 0.4
    for i, (label, color) in enumerate(stages):
        x = i * (box_w + 0.18)
        rect = plt.Rectangle((x, y), box_w, box_h,
                             facecolor=color, edgecolor="black",
                             linewidth=0.7, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + box_w / 2, y + box_h / 2, label,
                ha="center", va="center", fontsize=8.5,
                fontweight="bold", color="white")
        if i < n - 1:
            ax.annotate("", xy=(x + box_w + 0.16, y + box_h / 2),
                        xytext=(x + box_w + 0.02, y + box_h / 2),
                        arrowprops=dict(arrowstyle="->", lw=1.0))

    # Etichetta di output collegata all'ultimo stadio
    last_x = (n - 1) * (box_w + 0.18) + box_w
    ax.annotate(
        "BLOCK / CHALLENGE / ALLOW",
        xy=(last_x + 0.05, y + box_h / 2),
        xytext=(last_x + 0.6, y + box_h / 2 - 0.55),
        fontsize=8.5, fontweight="bold", color="#444",
        ha="left", va="center",
        arrowprops=dict(arrowstyle="-", lw=0.7, color="#666"),
    )

    ax.set_xlim(-0.2, n * (box_w + 0.18) + 1.6)
    ax.set_ylim(-0.2, 1.5)
    ax.axis("off")
    fig.suptitle("Pipeline di elaborazione di /api/bitm/collect",
                 fontsize=11, fontweight="bold", y=0.98)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_two_stage_scoring():
    out = FIGDIR / "two_stage.png"
    fig, ax = plt.subplots(figsize=(10, 3.6))

    # Tre box principali: Extractor | LLM Scorer | Policy
    boxes = [
        (0.5, "Extractor",
         "calcola pre_risk_score\ne confirmed_signals",
         "#F6BD16"),
        (3.7, "LLM Scorer",
         "ritorna risk_score\nverdict, indicators",
         "#9270CA"),
        (6.9, "Policy",
         "score = max(pre, llm)\n+ boost contestuale",
         "#E8684A"),
    ]
    for x, title, body, color in boxes:
        rect = plt.Rectangle((x, 1.2), 2.2, 1.8,
                             facecolor=color, edgecolor="black",
                             linewidth=0.7, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + 1.1, 2.65, title,
                ha="center", va="center", fontsize=11,
                fontweight="bold", color="white")
        ax.text(x + 1.1, 1.75, body,
                ha="center", va="center", fontsize=8.5,
                color="white")

    # Frecce orizzontali fra i box
    ax.annotate("", xy=(3.65, 2.1), xytext=(2.75, 2.1),
                arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.text(3.20, 2.30, "feature\n+ pre_score", ha="center",
            fontsize=8, color="#333")

    ax.annotate("", xy=(6.85, 2.1), xytext=(5.95, 2.1),
                arrowprops=dict(arrowstyle="->", lw=1.2))
    ax.text(6.40, 2.30, "risk_score", ha="center",
            fontsize=8, color="#333")

    # Box di output "Action"
    ax_box_x = 9.65
    rect = plt.Rectangle((ax_box_x, 1.55), 1.4, 1.1,
                         facecolor="#5AD8A6", edgecolor="black",
                         linewidth=0.7, alpha=0.9)
    ax.add_patch(rect)
    ax.text(ax_box_x + 0.7, 2.10, "Action",
            ha="center", va="center", fontsize=10,
            fontweight="bold", color="white")
    ax.text(ax_box_x + 0.7, 1.75,
            "allow / challenge\n/ block",
            ha="center", va="center", fontsize=7.5, color="white")
    ax.annotate("", xy=(ax_box_x - 0.05, 2.1), xytext=(9.15, 2.1),
                arrowprops=dict(arrowstyle="->", lw=1.2))

    # Arco rosso "floor": parte sotto Extractor e arriva sotto Policy,
    # passando SOTTO il box LLM Scorer per non sovrapporsi
    ax.annotate("", xy=(7.95, 1.18),
                xytext=(1.65, 1.18),
                arrowprops=dict(arrowstyle="->", lw=1.3, color="#C00",
                                connectionstyle="arc3,rad=0.35"))
    ax.text(4.85, -0.30, "pre_risk_score (floor)",
            ha="center", fontsize=9, color="#C00", fontweight="bold")

    ax.set_xlim(0.2, 11.2)
    ax.set_ylim(-0.5, 3.5)
    ax.axis("off")
    fig.suptitle("Modello di scoring a due stadi con floor deterministico",
                 fontsize=11, fontweight="bold", y=0.98)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_thresholds():
    out = FIGDIR / "thresholds.png"
    contexts = ["static", "default", "login", "admin", "payment"]
    challenge = [0.70, 0.40, 0.28, 0.22, 0.20]
    block     = [0.92, 0.75, 0.62, 0.60, 0.55]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    x = np.arange(len(contexts))

    # Allow zone (0 → challenge)
    for i in range(len(contexts)):
        ax.barh(i, challenge[i], color="#5AD8A6", alpha=0.85, height=0.6)
        ax.barh(i, block[i] - challenge[i], left=challenge[i],
                color="#F6BD16", alpha=0.85, height=0.6)
        ax.barh(i, 1 - block[i], left=block[i],
                color="#E8684A", alpha=0.85, height=0.6)
        # Numerical labels
        ax.text(challenge[i] - 0.01, i, f"{challenge[i]:.2f}",
                ha="right", va="center", fontsize=8.5,
                color="white", fontweight="bold")
        ax.text(block[i] - 0.01, i, f"{block[i]:.2f}",
                ha="right", va="center", fontsize=8.5,
                color="white", fontweight="bold")

    ax.set_yticks(x)
    ax.set_yticklabels(contexts)
    ax.set_xlim(0, 1.0)
    ax.set_xlabel("risk_score (0 = legittimo, 1 = attacco)")
    ax.set_title("Soglie contestuali (CHALLENGE, BLOCK)",
                 fontweight="bold")

    # Legend collocata sopra l'area dei dati per non coprire le barre
    from matplotlib.patches import Patch
    legend_elems = [
        Patch(facecolor="#5AD8A6", alpha=0.85, label="ALLOW"),
        Patch(facecolor="#F6BD16", alpha=0.85, label="CHALLENGE"),
        Patch(facecolor="#E8684A", alpha=0.85, label="BLOCK"),
    ]
    ax.legend(handles=legend_elems, loc="upper center",
              bbox_to_anchor=(0.5, -0.18), ncol=3,
              frameon=False, fontsize=9)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_signal_weights():
    out = FIGDIR / "weights.png"
    signals = [
        ("novnc_client_marker", 0.80),
        ("guacamole_client_marker", 0.80),
        ("bitm_framework_ua", 0.80),
        ("bitm_backend_port", 0.78),
        ("xss_reflected_param", 0.70),
        ("webauthn_api_override", 0.70),
        ("phantomjs_ua", 0.55),
        ("bitm_websocket_transport", 0.55),
        ("headlesschrome_ua", 0.50),
        ("webdriver_true", 0.45),
        ("extreme_latency", 0.35),
        ("tor_exit_node", 0.30),
        ("swiftshader_webgl", 0.30),
        ("tunnel_host", 0.25),
        ("empty_canvas", 0.15),
        ("no_languages", 0.15),
        ("iframe_overlay", 0.15),
        ("vpn_detected", 0.12),
        ("no_webgl_renderer", 0.12),
        ("no_timezone", 0.10),
    ]
    signals.sort(key=lambda x: x[1])
    labels = [s[0] for s in signals]
    weights = [s[1] for s in signals]

    fig, ax = plt.subplots(figsize=(9, 7))
    colors = ["#E8684A" if w >= 0.45 else
              "#F6BD16" if w >= 0.20 else
              "#5AD8A6" for w in weights]
    ax.barh(labels, weights, color=colors, alpha=0.85, edgecolor="black",
            linewidth=0.4)
    for i, w in enumerate(weights):
        ax.text(w + 0.01, i, f"{w:.2f}", va="center", fontsize=8)
    ax.set_xlabel("Peso nel pre_risk_score")
    ax.set_xlim(0, 1.0)
    ax.set_title("Pesi dei segnali deterministici (extractor._pre_score)",
                 fontweight="bold")

    from matplotlib.patches import Patch
    legend_elems = [
        Patch(facecolor="#E8684A", alpha=0.85, label="critico (≥0.45)"),
        Patch(facecolor="#F6BD16", alpha=0.85, label="medio (0.20–0.44)"),
        Patch(facecolor="#5AD8A6", alpha=0.85, label="debole (<0.20)"),
    ]
    ax.legend(handles=legend_elems, loc="lower right", frameon=False,
              fontsize=9)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_test_results():
    out = FIGDIR / "test_results.png"
    cats = ["legit", "attack", "suspicious", "edge", "system"]
    n_cases = [5, 13, 6, 5, 20]
    n_passed = [5, 13, 6, 5, 20]

    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(cats))
    width = 0.35
    ax.bar(x - width / 2, n_cases, width, label="Casi totali",
           color="#5B8FF9", alpha=0.85, edgecolor="black", linewidth=0.5)
    ax.bar(x + width / 2, n_passed, width, label="Casi superati",
           color="#5AD8A6", alpha=0.85, edgecolor="black", linewidth=0.5)
    for i, (t, p) in enumerate(zip(n_cases, n_passed)):
        ax.text(i - width / 2, t + 0.3, str(t), ha="center", fontsize=9)
        ax.text(i + width / 2, p + 0.3, str(p), ha="center", fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels(cats)
    ax.set_ylabel("Numero di casi")
    ax.set_title("Esito della suite di test per categoria (49/49 superati)",
                 fontweight="bold")
    ax.legend(frameon=False, fontsize=9)
    ax.set_ylim(0, max(n_cases) + 3)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_latency_distribution():
    out = FIGDIR / "latency.png"
    # Dati estratti da test_report.json
    fast_path = [4, 2, 2, 2, 2, 2, 3, 2, 2, 2, 2, 3, 3, 5, 3, 2]
    llm_path = [1402, 1290, 1284, 1656, 1638, 1604, 1593, 1774,
                1858, 1305, 1344, 1273, 1293]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    bp = ax.boxplot([fast_path, llm_path],
                    tick_labels=["Fast-track\ndeterministico\n(n=16)",
                                 "Pipeline LLM\ncompleta\n(n=13)"],
                    patch_artist=True, widths=0.5,
                    medianprops=dict(color="black", linewidth=1.2))
    bp["boxes"][0].set_facecolor("#5AD8A6")
    bp["boxes"][0].set_alpha(0.85)
    bp["boxes"][1].set_facecolor("#9270CA")
    bp["boxes"][1].set_alpha(0.85)
    ax.set_yscale("log")
    ax.set_ylabel("Latenza totale (ms, scala logaritmica)")
    ax.set_title("Distribuzione delle latenze osservate sulla test-suite",
                 fontweight="bold")
    ax.axhline(2000, color="#C00", linestyle="--", linewidth=0.8,
               alpha=0.6)
    ax.text(0.55, 2150, "Target p95 < 2000 ms",
            color="#C00", fontsize=8.5, va="bottom", ha="left",
            fontweight="bold")
    ax.grid(axis="y", which="both", linestyle=":", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_backend_comparison():
    out = FIGDIR / "backends.png"
    backends = ["Stub", "Anthropic\nHaiku 4.5", "Ollama\nllama3.1 7B"]
    accuracy = [100.0, 100.0, 100.0]
    latency = [3, 600, 1500]
    cost = [0.0, 0.001, 0.0]

    fig, axes = plt.subplots(1, 3, figsize=(11, 4.2))

    colors = ["#5AD8A6", "#5B8FF9", "#9270CA"]

    axes[0].bar(backends, accuracy, color=colors, alpha=0.85,
                edgecolor="black", linewidth=0.5)
    axes[0].set_ylim(0, 110)
    axes[0].set_ylabel("Accuratezza (%)")
    axes[0].set_title("Accuratezza", fontweight="bold")
    for i, v in enumerate(accuracy):
        axes[0].text(i, v + 2, f"{v:.0f}%", ha="center", fontsize=9)

    axes[1].bar(backends, latency, color=colors, alpha=0.85,
                edgecolor="black", linewidth=0.5)
    axes[1].set_ylabel("Latenza media (ms)")
    axes[1].set_title("Latenza", fontweight="bold")
    axes[1].set_yscale("log")
    for i, v in enumerate(latency):
        axes[1].text(i, v * 1.3, f"{v} ms", ha="center", fontsize=9)

    axes[2].bar(backends, cost, color=colors, alpha=0.85,
                edgecolor="black", linewidth=0.5)
    axes[2].set_ylabel("Costo per richiesta (USD)")
    axes[2].set_title("Costo", fontweight="bold")
    axes[2].set_ylim(0, 0.0015)
    for i, v in enumerate(cost):
        label = "0" if v == 0 else f"${v:.4f}"
        axes[2].text(i, v + 0.00005, label, ha="center", fontsize=9)

    for ax in axes:
        for tick in ax.get_xticklabels():
            tick.set_fontsize(9)

    fig.suptitle("Confronto fra i tre backend LLM supportati",
                 fontweight="bold", fontsize=12, y=0.98)
    fig.tight_layout(rect=(0, 0.02, 1, 0.94))
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_score_distribution():
    out = FIGDIR / "scores.png"
    # Dati estratti da test_report.json (campi score per ciascun caso T01-T29)
    scores_per_action = {
        "ALLOW (allow)": [0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0, 0.0],
        "CHALLENGE": [0.28, 0.27, 0.50, 0.22, 0.39, 0.43],
        "BLOCK": [1.0, 1.0, 0.97, 0.97, 1.0, 0.97, 0.97, 1.0, 0.97,
                  0.97, 0.97, 0.97, 0.97, 0.97],
    }
    fig, ax = plt.subplots(figsize=(9, 4.5))
    colors = {"ALLOW (allow)": "#5AD8A6",
              "CHALLENGE": "#F6BD16", "BLOCK": "#E8684A"}
    positions = [1, 2, 3]
    data = list(scores_per_action.values())
    bp = ax.boxplot(data, positions=positions, widths=0.55,
                    patch_artist=True,
                    medianprops=dict(color="black", linewidth=1.2),
                    showfliers=True)
    for patch, label in zip(bp["boxes"], scores_per_action.keys()):
        patch.set_facecolor(colors[label])
        patch.set_alpha(0.85)
    ax.set_xticks(positions)
    ax.set_xticklabels(list(scores_per_action.keys()))
    ax.set_ylabel("risk_score finale")
    ax.set_ylim(-0.05, 1.1)
    ax.axhline(0.40, color="#888", linestyle=":", linewidth=0.8)
    ax.axhline(0.75, color="#888", linestyle=":", linewidth=0.8)
    # Etichette dentro l'area dati per evitare clipping a destra
    ax.text(0.55, 0.42, "soglia challenge (default)",
            fontsize=8, color="#666", va="bottom", ha="left", style="italic")
    ax.text(0.55, 0.77, "soglia block (default)",
            fontsize=8, color="#666", va="bottom", ha="left", style="italic")
    ax.set_xlim(0.4, 3.6)
    ax.set_title("Distribuzione di risk_score per azione decisa",
                 fontweight="bold")
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_trajectory_patterns():
    out = FIGDIR / "trajectory.png"
    patterns = ["normal_flow", "panic_password\n_change",
                "direct_admin\n_access", "rapid_navigation",
                "insufficient\n_history"]
    scores = [0.0, 0.55, 0.40, 0.28, 0.0]
    colors = ["#5AD8A6", "#E8684A", "#F6BD16", "#F6BD16", "#CCCCCC"]

    fig, ax = plt.subplots(figsize=(9, 4.5))
    bars = ax.bar(patterns, scores, color=colors, alpha=0.85,
                  edgecolor="black", linewidth=0.5)
    ax.axhline(0.21, color="#666", linestyle="--", linewidth=0.8)
    ax.axhline(0.51, color="#C00", linestyle="--", linewidth=0.8)
    # Label collocate nell'angolo in alto a sinistra (zona libera)
    # per non sovrapporsi alle barre
    ax.text(-0.45, 0.665, "soglia compromissione (0.51)",
            fontsize=8.5, color="#C00", va="bottom", ha="left",
            fontweight="bold", style="italic")
    ax.text(-0.45, 0.135, "soglia sospetto (0.21)",
            fontsize=8.5, color="#666", va="bottom", ha="left",
            style="italic")
    ax.set_ylim(0, 0.7)
    ax.set_ylabel("trajectory_score")
    ax.set_title("Pattern di traiettoria riconosciuti dal layer Trajectory",
                 fontweight="bold")
    for b, s in zip(bars, scores):
        ax.text(b.get_x() + b.get_width() / 2, s + 0.015,
                f"{s:.2f}", ha="center", fontsize=9)
    plt.setp(ax.get_xticklabels(), fontsize=8.5)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def fig_bitm_taxonomy():
    out = FIGDIR / "taxonomy.png"
    fig, ax = plt.subplots(figsize=(9, 4.8))

    # MitM (top)
    ax.add_patch(plt.Rectangle((3, 4), 3, 0.7, facecolor="#5B8FF9",
                                alpha=0.85, edgecolor="black"))
    ax.text(4.5, 4.35, "Man-in-the-Middle",
            ha="center", va="center", fontweight="bold",
            color="white", fontsize=10)

    # Children
    children = [
        (0, "Man-in-the-Browser",
         "Trojan locale\n(Zeus, EKO, Sinowal)", "#9270CA"),
        (3, "Browser-in-the-Middle",
         "RFB/RDP via WebSocket\n(noVNC, Guacamole)", "#E8684A"),
        (6, "BitM+",
         "ngrok + Puppeteer\n+ evilGet (WebAuthn bypass)", "#C72C56"),
    ]
    for x, title, sub, col in children:
        ax.add_patch(plt.Rectangle((x, 1.3), 3, 1.7,
                                    facecolor=col, alpha=0.85,
                                    edgecolor="black"))
        ax.text(x + 1.5, 2.5, title, ha="center",
                fontweight="bold", color="white", fontsize=10)
        ax.text(x + 1.5, 1.85, sub, ha="center",
                color="white", fontsize=8)
        ax.annotate("", xy=(x + 1.5, 3),
                    xytext=(4.5, 4),
                    arrowprops=dict(arrowstyle="->", lw=0.9, color="#333"))

    ax.text(0.5, 0.7,
            "Richiede malware locale",
            fontsize=8, color="#666", style="italic")
    ax.text(3.5, 0.7,
            "No malware. Phishing + browser remoto",
            fontsize=8, color="#666", style="italic")
    ax.text(6.5, 0.7,
            "No malware + bypass MFA forte",
            fontsize=8, color="#666", style="italic")

    ax.set_xlim(-0.3, 9.5)
    ax.set_ylim(0, 5)
    ax.axis("off")
    fig.suptitle("Tassonomia degli attacchi browser-side",
                 fontweight="bold", fontsize=11)
    fig.tight_layout()
    fig.savefig(out)
    plt.close(fig)
    return out


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # Sfondo grigio chiaro tramite XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def init_doc_from_template():
    """Carica Template_Tesi.docx e svuota corpo, header e footer
    mantenendo stili, sezioni e impostazioni di pagina."""
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(
            f"Template non trovato: {TEMPLATE_PATH}. "
            "Posizionare 'Template_Tesi.docx' nella radice del progetto."
        )
    doc = Document(str(TEMPLATE_PATH))

    body = doc.element.body
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)

    # Pulisci header/footer del template (riferiscono alla vecchia tesi)
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for p in list(container.paragraphs):
                p._element.getparent().remove(p._element)
            container.add_paragraph()

    return doc


def add_capitolo(doc, numero_romano, titolo):
    """Stile capitoli del template: Heading 1 'Capitolo N' + Subtitle col titolo."""
    h = doc.add_heading(f"Capitolo {numero_romano}", level=1)
    try:
        sub = doc.add_paragraph(titolo, style="Subtitle")
    except KeyError:
        sub = doc.add_paragraph(titolo)
        for r in sub.runs:
            r.bold = True
            r.font.size = Pt(18)
    return h, sub


def setup_styles(doc):
    """Il template definisce già tutti gli stili. Lasciamo intatti i Heading."""
    return


def page_setup(doc):
    """Margini ed impostazioni di pagina arrivano dal template."""
    return


TITOLO_TESI = (
    "Rilevamento di attacchi Browser-in-the-Middle tramite Large "
    "Language Models: progettazione e sviluppo di un'architettura "
    "di difesa in tempo reale"
)

NOME_LAUREANDO = "Gabriele SANZIONE"
NOME_RELATORE = "Chiar.mo Prof. Danilo CAIVANO"
NOME_CORRELATORE = "Dott.ssa Vita Santa BARLETTA"
ANNO_ACCADEMICO = "2025/2026"
SUBJECT_TESI = "SICUREZZA INFORMATICA"


def _para_centered(doc, runs, space_after=0, line_spacing=None):
    """Helper interno: paragrafo centrato con run multipli (text, size_pt, bold, italic)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    for spec in runs:
        text = spec["text"]
        size = spec.get("size", 14)
        bold = spec.get("bold", False)
        italic = spec.get("italic", False)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
    return p


def frontespizio(doc):
    # Spaziatura iniziale
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    _para_centered(doc, [
        {"text": "Università degli Studi di Bari ", "size": 26},
        {"text": "“", "size": 26},
        {"text": "Aldo Moro", "size": 26},
        {"text": "”", "size": 26},
    ], space_after=4)

    _para_centered(doc, [
        {"text": "Dipartimento di Informatica", "size": 14},
    ], space_after=2)
    _para_centered(doc, [
        {"text": "Corso di Laurea in Informatica", "size": 14},
    ], space_after=24)

    _para_centered(doc, [
        {"text": "Tesi di laurea in", "size": 14},
    ], space_after=4)
    _para_centered(doc, [
        {"text": SUBJECT_TESI, "size": 14, "bold": True},
    ], space_after=18)

    # Titolo della tesi nello stile "Sottotitolo Documento" del template
    try:
        ptit = doc.add_paragraph(style="Sottotitolo Documento")
    except KeyError:
        ptit = doc.add_paragraph()
    ptit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ptit.paragraph_format.space_before = Pt(12)
    ptit.paragraph_format.space_after = Pt(36)
    rtit = ptit.add_run(TITOLO_TESI)
    rtit.font.name = "Times New Roman"
    rtit.font.size = Pt(22)

    # Relatore (allineato a sinistra)
    p_rel_lab = doc.add_paragraph()
    p_rel_lab.paragraph_format.space_after = Pt(0)
    r = p_rel_lab.add_run("Relatore:")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.space_after = Pt(0)
    r = p_rel.add_run(NOME_RELATORE)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True

    p_cor = doc.add_paragraph()
    p_cor.paragraph_format.space_after = Pt(18)
    r = p_cor.add_run(NOME_CORRELATORE)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True

    # Laureando (allineato a destra)
    p_lau_lab = doc.add_paragraph()
    p_lau_lab.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_lau_lab.paragraph_format.space_after = Pt(0)
    r = p_lau_lab.add_run("Laureando:")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    p_lau = doc.add_paragraph()
    p_lau.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_lau.paragraph_format.space_after = Pt(36)
    r = p_lau.add_run(NOME_LAUREANDO)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True

    # Anno accademico
    _para_centered(doc, [
        {"text": "Anno Accademico ", "size": 16},
        {"text": ANNO_ACCADEMICO, "size": 16},
    ], space_after=0)

    add_page_break(doc)


def setup_header_footer(doc):
    """Imposta header e footer come da template ma con il titolo della tesi corrente."""
    sec = doc.sections[0]

    # Header: tab + "Indice"  (verrà sostituito automaticamente da Word ai capitoli)
    h = sec.header
    if not h.paragraphs:
        h.add_paragraph()
    hp = h.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("\t")
    hr.font.name = "Times New Roman"
    hr.font.size = Pt(10)
    hr2 = hp.add_run("Tesi di Laurea")
    hr2.italic = True
    hr2.font.name = "Times New Roman"
    hr2.font.size = Pt(10)

    # Footer: titolo tesi a sinistra + numero pagina a destra (campo PAGE)
    f = sec.footer
    if not f.paragraphs:
        f.add_paragraph()
    fp = f.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(
        Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )
    fr = fp.add_run(TITOLO_TESI[:80])
    fr.italic = True
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)
    fp.add_run("\t")

    # Campo PAGE (numero pagina dinamico)
    run_page = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_page._r.append(fld_begin)
    run_page._r.append(instr)
    run_page._r.append(fld_sep)
    run_page._r.append(txt)
    run_page._r.append(fld_end)
    run_page.font.name = "Times New Roman"
    run_page.font.size = Pt(9)


def dedica(doc):
    for _ in range(8):
        doc.add_paragraph()
    add_par(doc,
            "Ai miei genitori,\nche hanno sempre creduto in me\n"
            "anche quando io stesso vacillavo.",
            italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=False, space_after=4)
    add_page_break(doc)


def indice(doc):
    """Inserisce il Sommario come campo TOC di Word.
    All'apertura, l'utente può aggiornarlo con F9 (o click destro → Aggiorna campo).
    Il TOC raccoglie automaticamente tutti gli Heading 1/2/3 del documento."""
    add_heading(doc, "Indice", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")  # forza Word a chiedere l'aggiornamento

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "[Aggiornare il sommario premendo F9 in Word "
        "oppure click destro → Aggiorna campo]"
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)

    add_page_break(doc)


# ============================================================================
# CAPITOLI
# ============================================================================

def introduzione(doc):
    add_capitolo(doc, "I", "Introduzione")

    add_par(doc,
        "La superficie d'attacco della rete Internet si è andata ampliando in "
        "modo significativo negli ultimi anni, sospinta dall'aumento dei "
        "servizi web critici — banking, e-commerce, identità digitale — e "
        "dalla parallela industrializzazione delle tecniche di phishing. "
        "Fra le minacce emergenti, una posizione di particolare rilievo è "
        "occupata dal cosiddetto attacco Browser-in-the-Middle (BitM), "
        "formalizzato da Tommasi, Catalano e Taurino dell'Università del "
        "Salento nel 2021. Diversamente dal classico Man-in-the-Middle, "
        "che richiede il controllo del canale di comunicazione, e dal "
        "Man-in-the-Browser, che presuppone l'installazione di un trojan "
        "sulla macchina della vittima, il BitM realizza l'intercettazione "
        "interponendo un browser malevolo trasparente, ospitato sul server "
        "dell'attaccante e fruito dalla vittima attraverso il proprio "
        "browser legittimo per mezzo del protocollo RFB incapsulato in "
        "WebSocket.")

    add_par(doc,
        "L'evoluzione più recente, denominata BitM+ e descritta nel 2025 "
        "da Catalano e collaboratori, sostituisce lo stack RFB con una "
        "combinazione di tunnel HTTPS pubblico (tipicamente ngrok), "
        "browser headless controllato via Puppeteer e payload XSS riflessi, "
        "ottenendo come risultato il bypass effettivo dei meccanismi di "
        "autenticazione forte basati su WebAuthn/FIDO2. Questa nuova "
        "categoria di attacchi rappresenta una sfida di prima grandezza "
        "per le architetture di difesa, perché elude per costruzione tanto "
        "i sistemi anti-phishing basati su URL reputation quanto le "
        "soluzioni MFA tradizionali.")

    add_par(doc,
        "L'oggetto della presente tesi è la progettazione, lo sviluppo e "
        "la validazione sperimentale di un sistema di rilevamento "
        "real-time degli attacchi BitM e BitM+, denominato BitM Detection "
        "Plugin (di seguito anche \"BitM-LLM\"), che integra in un'unica "
        "pipeline difensiva tre paradigmi complementari: regole "
        "deterministiche a latenza zero per i segnali certi, scoring "
        "statistico contestuale per i segnali deboli e analisi tramite "
        "Large Language Model (LLM) per le valutazioni qualitative del "
        "contesto di sessione. Il sistema è realizzato in Python come "
        "servizio FastAPI, accompagnato da un'estensione browser MV3 "
        "(BitM Shield) e da una dashboard real-time WebSocket per il "
        "monitoraggio operativo.")

    add_par(doc,
        "Il presente lavoro è stato sviluppato presso il Dipartimento di "
        "Informatica dell'Università degli Studi di Bari Aldo Moro, e si "
        "inserisce nel filone di ricerca che i gruppi di Bari, Salento e "
        "altri atenei italiani conducono da diversi anni sul tema della "
        "sicurezza delle applicazioni web, in particolare con le "
        "pubblicazioni del 2021 sul BitM, del 2023 sull'educazione alla "
        "sicurezza tramite metafore narrative e del 2025 sull'analisi "
        "comparativa dei toolkit BitM.")

    add_heading(doc, "Motivazione e contributo originale", level=2)

    add_par(doc,
        "I sistemi di rilevamento attualmente diffusi nel mercato — Web "
        "Application Firewall, soluzioni di bot management, motori di "
        "fingerprinting — adottano in larga parte un approccio binario "
        "fondato su signature statiche o su modelli di machine learning "
        "addestrati offline su dataset proprietari. Tale approccio si "
        "scontra con due limiti strutturali: in primo luogo, i toolkit "
        "BitM — Evilginx, Modlishka, EvilProxy, e i nuovi framework "
        "BitM+ — vengono aggiornati con frequenza superiore al "
        "release-cycle medio dei modelli ML, generando una finestra di "
        "scopertura non trascurabile; in secondo luogo, molti segnali "
        "BitM — un titolo \"noVNC\" nella pagina, una porta TCP "
        "anomala nell'URL, una traiettoria di navigazione che salti il "
        "login per accedere direttamente a /admin — non sono "
        "categorizzabili a priori come maligni, ma richiedono una "
        "valutazione contestuale che è esattamente il dominio di "
        "competenza dei moderni LLM.")

    add_par(doc,
        "Il contributo originale di questo lavoro consiste in tre "
        "elementi principali. Primo, una pipeline ibrida che separa "
        "in modo netto i segnali a rischio certo (gestiti deterministicamente "
        "in meno di un millisecondo) dai segnali ambigui (delegati al LLM), "
        "ottenendo come effetto collaterale una drastica riduzione del "
        "costo di esercizio: i casi T01–T11 della suite di test, che "
        "rappresentano il 22% del traffico di prova, vengono risolti senza "
        "neppure invocare l'LLM. Secondo, l'introduzione di un secondo "
        "layer LLM dedicato all'analisi della traiettoria di sessione, "
        "capace di riconoscere pattern post-compromissione come il "
        "\"panic password change\" — login seguito da cambio password "
        "in meno di cinque secondi — che il solo fingerprint non può "
        "rivelare. Terzo, un'architettura modulare a tre backend LLM "
        "intercambiabili (Anthropic Claude, Ollama llama3.1 locale, "
        "stub deterministico) che consente di calibrare il trade-off fra "
        "qualità delle decisioni, costo per richiesta e privacy dei dati.")

    add_heading(doc, "Struttura della tesi", level=2)

    add_par(doc,
        "L'elaborato è organizzato in sei capitoli più la presente "
        "introduzione. Il primo capitolo presenta lo stato dell'arte: "
        "ricostruisce la genesi tecnica dell'attacco BitM, ne illustra "
        "l'evoluzione verso BitM+ e descrive le contromisure proposte "
        "in letteratura, evidenziandone i limiti rispetto alla minaccia "
        "attuale; chiude con una panoramica sulle applicazioni dei "
        "Large Language Models nel campo della cybersecurity.")

    add_par(doc,
        "Il secondo capitolo formalizza il problema affrontato, definisce "
        "obiettivi, requisiti funzionali e non funzionali e documenta i "
        "vincoli adottati. Il terzo capitolo descrive l'architettura del "
        "sistema, motivando le scelte di progettazione — pipeline a "
        "stadi, scoring a due livelli, soglie contestuali, layer "
        "trajectory — alla luce dei requisiti del capitolo precedente.")

    add_par(doc,
        "Il quarto capitolo entra nel merito implementativo modulo per "
        "modulo, riportando le scelte chiave e gli snippet di codice più "
        "rilevanti. Il quinto capitolo presenta il piano di "
        "sperimentazione, articolato in 49 casi di test che coprono "
        "scenari legittimi, attacchi conclamati, casi sospetti, "
        "edge-case e controlli di sistema, e ne discute i risultati. "
        "Il capitolo finale tira le somme del lavoro svolto e indica le "
        "direzioni di ricerca futura, con particolare attenzione alle "
        "possibilità di fine-tuning specializzato del modello "
        "Llama 3.1 mediante tecniche LoRA, già infrastrutturate nel "
        "repository ma non ancora oggetto di valutazione su larga scala.")

    add_par(doc,
        "Una bibliografia commentata e un breve elenco di "
        "ringraziamenti chiudono il documento.")

    add_page_break(doc)


def capitolo1_stato_arte(doc):
    add_capitolo(doc, "II", "Stato dell'arte")

    add_heading(doc, "1.1 Dal Man-in-the-Middle al Browser-in-the-Middle",
                level=2)

    add_par(doc,
        "L'attacco Man-in-the-Middle (MitM) è uno dei vettori di "
        "compromissione più studiati nella letteratura di sicurezza "
        "informatica. Nella sua formulazione canonica, l'attaccante si "
        "interpone in un punto del canale di comunicazione fra due "
        "endpoint legittimi e ne intercetta, modifica o blocca il "
        "traffico. La compromissione può colpire una qualsiasi delle "
        "tre proprietà fondamentali della sicurezza: la riservatezza "
        "(eavesdropping), l'integrità (manipolazione del payload), la "
        "disponibilità (interruzione o sabotaggio della comunicazione). "
        "L'ampia varietà di tecnologie attaccabili — LTE, Bluetooth, "
        "NFC, HTTPS, WiFi, processi di sistema operativo — testimonia "
        "la versatilità di questa categoria.")

    add_par(doc,
        "Il principale ostacolo all'esecuzione di un MitM è di natura "
        "operativa: l'attaccante deve riuscire ad acquisire un punto "
        "di accesso al canale, sia esso fisico (sniffing su una rete "
        "WiFi pubblica) o logico (sfruttamento di una vulnerabilità "
        "zero-day, di un certificato compromesso, di un router "
        "configurato in modo errato). Come conseguenza, gli attacchi "
        "MitM \"puri\" hanno una scalabilità limitata e spesso "
        "richiedono la prossimità fisica della vittima.")

    add_par(doc,
        "Il Man-in-the-Browser (MitB) supera il vincolo di prossimità "
        "spostando l'intercettazione nel browser stesso della vittima, "
        "tipicamente attraverso un trojan installato in modo "
        "fraudolento (Zeus, Adrenaline, Sinowal, Silent Banker, "
        "Eurograbber, EKO sono fra gli esempi più noti documentati in "
        "letteratura). Tuttavia il MitB conserva un punto debole: "
        "richiede l'esecuzione di codice nativo sulla macchina target, "
        "un'azione che diventa progressivamente più difficile mano a "
        "mano che i sistemi operativi rinforzano i propri meccanismi "
        "anti-malware (Windows Defender, SmartScreen, AppLocker, "
        "controllo applicazioni in macOS).")

    add_par(doc,
        "Il Browser-in-the-Middle (BitM), formalizzato da Tommasi, "
        "Catalano e Taurino nel paper pubblicato sull'International "
        "Journal of Information Security nel 2021, supera entrambi i "
        "vincoli precedenti. La sua idea di fondo è radicalmente "
        "diversa: invece di aggredire il canale o il browser della "
        "vittima, sposta interamente il browser sul server "
        "dell'attaccante. Il browser \"trasparente\" dell'attaccante "
        "instaura una connessione legittima con il sito target e ne "
        "espone l'interfaccia visiva al browser della vittima "
        "attraverso il protocollo RFB (Remote Frame Buffer) "
        "veicolato su WebSocket via JavaScript.")

    add_figure(doc, fig_bitm_taxonomy(),
               "Figura 1.1 — Tassonomia degli attacchi browser-side: dal MitM "
               "al Browser-in-the-Middle e alla sua evoluzione BitM+.")

    add_par(doc,
        "Da un punto di vista percettivo, la vittima non vede "
        "nulla di anomalo: il sito appare identico all'originale, "
        "perché lo è — è realmente l'originale, ma navigato per "
        "interposta persona. Da un punto di vista difensivo, "
        "l'attaccante ha pieno controllo del flusso di dati: vede "
        "le credenziali digitate, accede al cookie di sessione, "
        "può manipolare il payload, può continuare la sessione "
        "dopo che la vittima ha chiuso il browser. La principale "
        "innovazione difensiva sulla quale si poggia il BitM è "
        "l'eliminazione della necessità di malware locale: tutto "
        "ciò che la vittima deve fare è cliccare un link.")

    add_heading(doc, "1.2 Anatomia tecnica dell'attacco BitM", level=2)

    add_par(doc,
        "Il prototipo descritto in Tommasi 2021 si fonda su una "
        "piattaforma GNU/Linux equipaggiata con un window manager "
        "minimale (Fluxbox), un browser Chromium customizzato per "
        "rimuovere ogni elemento di chrome (frame, barra degli "
        "indirizzi, schede), un server VNC e il pacchetto noVNC. "
        "Il tutto è esposto al pubblico tramite un web server HTTP "
        "che serve la pagina HTML5 contenente il client JavaScript "
        "noVNC.")

    add_par(doc,
        "Il flusso di esecuzione dell'attacco si articola nei "
        "seguenti passi:")

    add_bullet(doc, "L'attaccante predispone un link malevolo che "
                    "punta al proprio server BitM. Il link viene "
                    "diffuso tramite phishing (email, SMS, social).")
    add_bullet(doc, "La vittima clicca sul link e il proprio "
                    "browser carica la pagina HTML5 servita dal "
                    "server BitM.")
    add_bullet(doc, "La pagina HTML5 contiene il client noVNC, un "
                    "programma JavaScript che apre una connessione "
                    "WebSocket verso il server malevolo.")
    add_bullet(doc, "Il componente Websockify, in esecuzione sul "
                    "server, traduce il traffico WebSocket in "
                    "traffico VNC binario verso il vncserver locale.")
    add_bullet(doc, "Il vncserver locale renderizza graficamente "
                    "il browser Chromium customizzato, che a sua "
                    "volta è connesso al sito target legittimo.")
    add_bullet(doc, "La vittima vede sullo schermo il sito legittimo, "
                    "ma sta in realtà guardando un \"video stream\" "
                    "delle azioni del browser dell'attaccante.")
    add_bullet(doc, "Ogni azione che la vittima compie (movimento "
                    "mouse, pressione tasti) viene trasmessa al server "
                    "BitM, che la inoltra al proprio browser, il quale "
                    "la propaga al sito target.")

    add_par(doc,
        "Una variante alternativa, descritta nello stesso paper "
        "Tommasi 2021 e consolidata da Tzschoppe nel 2023, sostituisce "
        "lo stack RFB/noVNC con il protocollo RDP (Remote Desktop "
        "Protocol) attraverso il gateway open-source Apache Guacamole. "
        "Guacamole offre supporto per RDP, VNC e SSH attraverso un "
        "client HTML5 e funziona come proxy tra il browser HTML5 "
        "client e una connessione RDP/VNC server-side. Dal punto di "
        "vista dell'attaccante, la variante Guacamole è interessante "
        "perché RDP supporta nativamente la condivisione del clipboard "
        "e dei dispositivi USB, ampliando ulteriormente il perimetro "
        "di compromissione.")

    add_heading(doc, "1.3 BitM+: phishing-as-a-service e bypass di WebAuthn",
                level=2)

    add_par(doc,
        "Il paper di Catalano e collaboratori del 2025 introduce un "
        "salto qualitativo nella minaccia: la categoria denominata "
        "BitM+ rappresenta un'industrializzazione del BitM in chiave "
        "phishing-as-a-service e include come obiettivo esplicito "
        "il bypass dei sistemi di autenticazione forte basati su "
        "WebAuthn/FIDO2.")

    add_par(doc,
        "L'architettura tipica di un BitM+ comprende quattro "
        "componenti:")

    add_bullet(doc, "Un tunnel HTTPS pubblico, generalmente ottenuto "
                    "tramite servizi come ngrok, trycloudflare o "
                    "localtunnel, che fornisce all'attaccante un "
                    "certificato TLS valido senza necessità di "
                    "domini o ACME. Questo aspetto è cruciale, "
                    "perché WebAuthn richiede il dominio originale "
                    "per validare l'asserzione: senza tunnel HTTPS "
                    "credibile, il BitM+ non funziona.")
    add_bullet(doc, "Un browser headless controllato da Puppeteer, "
                    "che impersona la vittima sul sito target. La "
                    "scelta di Puppeteer (su Chromium) anziché di "
                    "Selenium è motivata dalla migliore performance "
                    "in WebSocket throughput e dal supporto nativo "
                    "per la simulazione di eventi touch.")
    add_bullet(doc, "Un server Node/Express, denominato MalSrv nei "
                    "papers, in ascolto sulla porta 3081, che orchestra "
                    "la sessione BitM lato server, riceve gli input "
                    "dal client e inoltra le risposte.")
    add_bullet(doc, "Un payload XSS riflesso, denominato evilGet(), "
                    "che sostituisce dinamicamente l'implementazione "
                    "nativa di navigator.credentials.get() del browser "
                    "vittima per intercettare la richiesta WebAuthn "
                    "e dirigerla verso il browser headless.")

    add_par(doc,
        "L'aspetto più insidioso di BitM+ è che riesce a effettuare "
        "il bypass di WebAuthn — pubblicamente raccomandato come "
        "stato dell'arte dell'autenticazione resistente al phishing — "
        "grazie a una combinazione di ingegneria che sfrutta "
        "esattamente il margine di confidenza che gli utenti hanno "
        "verso i tunnel HTTPS \"validi\". L'utente che si vede "
        "presentare un certificato TLS verde su un dominio "
        "ngrok-free.app non sa, e di norma non ha gli strumenti "
        "cognitivi per verificare, che quel certificato non è "
        "quello del sito di destinazione.")

    add_heading(doc, "1.4 Approcci difensivi esistenti e loro limiti", level=2)

    add_par(doc,
        "Le contromisure proposte in letteratura per il BitM si "
        "possono raggruppare in tre categorie principali: "
        "URL-based, fingerprinting-based e behavior-based. "
        "Ciascuna presenta limiti specifici che motivano il "
        "ricorso a un approccio integrato.")

    add_par(doc,
        "Il primo approccio, URL-based, è esemplificato dal "
        "lavoro di Catalano e collaboratori del 2023 che propone "
        "una mitigation basata sul controllo in tempo reale degli "
        "URL navigati confrontati contro il servizio Norton "
        "Safe Web. La logica è semplice: appena l'utente apre un "
        "URL sconosciuto, il sistema scarica la valutazione "
        "(Secure / Untested / Dangerous / Caution) e, in caso "
        "negativo, aggiunge il dominio alla blacklist hosts. "
        "Il limite è duplice: da un lato la latenza della "
        "verifica esterna, dall'altro la natura by-design di "
        "BitM+ che adopera domini ngrok freschi, mai indicizzati "
        "da feed di reputation.")

    add_par(doc,
        "Il secondo approccio, fingerprinting-based, è "
        "incarnato dai prodotti commerciali di bot management "
        "(Akamai Bot Manager, Cloudflare Turnstile, DataDome). "
        "Funziona estraendo decine di feature dal browser — "
        "User-Agent, plugin, WebGL renderer, canvas hash, "
        "timezone — e sottoponendole a un classificatore "
        "(spesso un Gradient Boosted Tree) addestrato su "
        "traffici reali etichettati. Il punto debole, "
        "documentato in letteratura, è la rapidità con cui "
        "i toolkit BitM/BitM+ rilasciano patch che neutralizzano "
        "le signature: il release-cycle medio del modello "
        "ML è dell'ordine delle settimane, mentre il "
        "release-cycle dei toolkit è dell'ordine dei giorni.")

    add_par(doc,
        "Il terzo approccio, behavior-based, è quello adottato "
        "dai web application firewall di nuova generazione e "
        "dai sistemi UEBA (User and Entity Behavior Analytics). "
        "Si basa sull'osservazione che gli attacchi automatizzati "
        "lasciano firme caratteristiche nella sequenza temporale "
        "delle richieste — burst di scraping, accessi a endpoint "
        "amministrativi senza preventiva autenticazione, cambi "
        "di password subito dopo il login. Il limite di questi "
        "sistemi è la difficoltà di definire a priori l'insieme "
        "completo di pattern: una regola scritta a mano è troppo "
        "rigida, un modello statistico richiede grandi volumi "
        "di traffico etichettato.")

    add_par(doc,
        "Nessuno dei tre approcci, isolatamente, è in grado di "
        "garantire una copertura adeguata contro la generazione "
        "attuale di toolkit BitM+, e l'ipotesi di lavoro che "
        "guida questa tesi è che l'integrazione dei tre — "
        "moderata da un Large Language Model che svolga la "
        "funzione di valutatore qualitativo del contesto — "
        "possa offrire un livello di protezione superiore.")

    add_heading(doc, "1.5 Large Language Models nella sicurezza informatica",
                level=2)

    add_par(doc,
        "L'applicazione dei Large Language Models a problemi di "
        "sicurezza informatica è un campo di ricerca in rapida "
        "espansione. I principali utilizzi documentati in "
        "letteratura riguardano il riassunto di alert SIEM, la "
        "generazione automatica di regole Sigma o YARA, l'analisi "
        "di campioni di malware tramite reverse engineering "
        "assistito, la triage di vulnerabilità in pipeline DevOps "
        "e — più di recente — la classificazione di richieste "
        "HTTP sospette in tempo reale.")

    add_par(doc,
        "Nel contesto specifico della rilevazione di attacchi "
        "browser-side, l'LLM si presta come componente di "
        "valutazione contestuale per tre ragioni. In primo luogo, "
        "i segnali di un attacco BitM sono raramente univoci "
        "presi singolarmente: una latenza di 350 ms su una "
        "pagina di pagamento non è di per sé indicativa, ma "
        "diventa sospetta se si combina con un canvas vuoto e "
        "una VPN. La capacità del LLM di pesare combinazioni "
        "non lineari di indicatori senza richiedere un re-training "
        "ad-hoc è un vantaggio rispetto ai classificatori "
        "tradizionali.")

    add_par(doc,
        "In secondo luogo, l'LLM è in grado di generare "
        "spiegazioni in linguaggio naturale del proprio verdetto, "
        "una caratteristica che si rivela preziosa sia per "
        "l'operatore di sicurezza che deve decidere come "
        "rispondere a un allarme, sia per l'utente finale che "
        "ha diritto di sapere perché una sua azione è stata "
        "bloccata. Nel sistema sviluppato in questa tesi, le "
        "spiegazioni in italiano destinate all'utente finale "
        "sono prodotte direttamente dal layer Trajectory.")

    add_par(doc,
        "In terzo luogo, l'LLM consente di sostituire una parte "
        "consistente del lavoro di feature engineering con un "
        "prompt ben costruito. Ciò non significa rinunciare alle "
        "feature deterministiche — anzi, il sistema descritto "
        "ne calcola decine — ma demandare al LLM la sintesi "
        "qualitativa dei segnali, alleggerendo lo sforzo di "
        "manutenzione del codice.")

    add_par(doc,
        "Il principale rischio dell'approccio è il fenomeno "
        "dell'allucinazione: un LLM può occasionalmente produrre "
        "verdetti incoerenti con i propri input, o assegnare "
        "una probabilità di attacco a un browser palesemente "
        "legittimo. La strategia di mitigazione adottata nel "
        "sistema BitM-LLM è il floor sul pre_risk_score "
        "deterministico: l'LLM non può scendere sotto la "
        "valutazione minima derivata da segnali certi, e in "
        "presenza di indicatori critici (CRITICAL_BLOCK) il "
        "verdetto LLM viene bypassato del tutto. Questo "
        "garantisce che il modello, anche nel caso peggiore, "
        "non possa peggiorare le decisioni rispetto al "
        "baseline deterministico.")

    add_page_break(doc)


def capitolo2_problema(doc):
    add_capitolo(doc, "III", "Analisi del problema e requisiti")

    add_heading(doc, "2.1 Obiettivi del progetto", level=2)

    add_par(doc,
        "L'obiettivo primario del progetto BitM-LLM è la "
        "realizzazione di un sistema di rilevamento real-time "
        "degli attacchi Browser-in-the-Middle che soddisfi "
        "contemporaneamente quattro requisiti di carattere "
        "qualitativo: efficacia rispetto allo stato dell'arte "
        "degli attacchi noti, latenza compatibile con "
        "l'integrazione in un percorso di richiesta HTTP "
        "produttivo, costi di esercizio sostenibili in scenari "
        "di traffico medio-alto, intercambiabilità del backend "
        "LLM per consentire deployment in scenari diversificati "
        "(cloud commerciale, ambiente self-hosted, ambiente "
        "air-gapped per pubblica amministrazione).")

    add_par(doc,
        "Più nello specifico, il sistema deve essere in grado "
        "di classificare ogni richiesta HTTP entrante in una "
        "delle tre categorie di azione: ALLOW (la richiesta "
        "viene servita normalmente), CHALLENGE (la richiesta "
        "viene sottoposta a una verifica aggiuntiva, "
        "tipicamente un CAPTCHA o una step-up di autenticazione), "
        "BLOCK (la richiesta viene rifiutata con codice HTTP "
        "appropriato e l'IP entra in escalation di blocco).")

    add_par(doc,
        "La copertura mira esplicitamente alle tre famiglie di "
        "minaccia documentate nei paper di riferimento: lo stack "
        "BitM RFB (noVNC + Websockify + TigerVNC), lo stack BitM "
        "RDP (Apache Guacamole + FreeRDP), e lo stack BitM+ "
        "(ngrok + Puppeteer + MalSrv + evilGet). Per ciascuna "
        "famiglia il sistema deve esporre segnali specifici, "
        "etichettati in modo univoco, che consentano "
        "l'attribuzione del verdetto al toolkit responsabile.")

    add_heading(doc, "2.2 Requisiti funzionali", level=2)

    rf = [
        ("RF-01", "Rilevamento di User-Agent automatizzati noti "
                  "(HeadlessChrome, PhantomJS, Selenium, Puppeteer, "
                  "JSDOM, SlimerJS)."),
        ("RF-02", "Riconoscimento del flag webdriver impostato "
                  "(navigator.webdriver === true)."),
        ("RF-03", "Identificazione dei marker BitM specifici: titolo "
                  "noVNC/Guacamole, porte di backend (3081/6080/4822/5900), "
                  "host di tunneling (ngrok, trycloudflare, localtunnel)."),
        ("RF-04", "Rilevamento del bypass WebAuthn tramite override "
                  "di navigator.credentials.get()."),
        ("RF-05", "Calcolo di un punteggio di rischio pre-LLM "
                  "deterministico, basato su pesi configurabili."),
        ("RF-06", "Invocazione del LLM per la valutazione "
                  "contestuale dei segnali ambigui."),
        ("RF-07", "Applicazione di soglie contestuali differenziate "
                  "per pagine login, payment, admin, default, static."),
        ("RF-08", "Persistenza delle sessioni utente con "
                  "tracciamento delle pagine visitate e dei tempi "
                  "di richiesta."),
        ("RF-09", "Rate-limiting per IP sorgente con sliding window."),
        ("RF-10", "Escalation di blocco IP dopo N rifiuti consecutivi."),
        ("RF-11", "Logging strutturato di ogni evento in formato "
                  "JSON Lines per analisi offline."),
        ("RF-12", "Notifica push verso Slack, Microsoft Teams o "
                  "endpoint SIEM in caso di evento BLOCK."),
        ("RF-13", "Dashboard real-time in tempo reale con feed "
                  "WebSocket degli eventi."),
        ("RF-14", "Endpoint di health-check che esponga lo stato "
                  "complessivo del sistema (backend, store, geoip, "
                  "webhook, sessioni attive)."),
        ("RF-15", "Layer di analisi della traiettoria di sessione "
                  "per il rilevamento di pattern post-compromissione "
                  "(panic password change, direct admin access, "
                  "rapid navigation)."),
        ("RF-16", "Estensione browser MV3 (Chrome/Edge/Brave) per "
                  "la protezione lato utente, con modalità local "
                  "(senza rete) e hybrid (con consultazione del "
                  "backend)."),
    ]

    for code, desc in rf:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        run1 = p.add_run(f"{code}. ")
        run1.bold = True
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(12)
        run2 = p.add_run(desc)
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(12)

    add_heading(doc, "2.3 Requisiti non funzionali", level=2)

    add_par(doc,
        "Sul fronte non funzionale, il sistema è stato "
        "progettato attorno ai seguenti obiettivi di qualità:")

    add_bullet(doc, "Latenza p95 della pipeline complessiva inferiore "
                    "a 2 secondi per richiesta (incluso il round-trip "
                    "verso l'LLM remoto). Per i casi gestiti dal fast-track "
                    "deterministico, latenza inferiore a 5 ms.")
    add_bullet(doc, "Capacità di servire almeno 30 richieste al minuto "
                    "per IP sorgente prima di scattare il rate-limit "
                    "(valore configurabile via RATE_LIMIT).")
    add_bullet(doc, "Disponibilità del 99.5% in modalità degradata: il "
                    "fallimento dell'LLM remoto o di Redis non deve "
                    "compromettere la funzionalità del servizio, che "
                    "ricade su backend stub e session store in-memory.")
    add_bullet(doc, "Portabilità: deploy via Docker Compose, "
                    "containerizzazione minimale, requisito singolo "
                    "Python 3.11+ per esecuzione bare-metal.")
    add_bullet(doc, "Configurabilità via file .env, senza necessità "
                    "di ricompilare o modificare il codice per cambiare "
                    "backend LLM, soglie, TTL di cache.")
    add_bullet(doc, "Sicurezza by-design: nessun secret hardcoded, "
                    "endpoint amministrativi protetti da token, "
                    "X-Forwarded-For accettato solo da proxy esplicitamente "
                    "fidati per evitare IP spoofing.")
    add_bullet(doc, "Privacy: in modalità local l'estensione browser "
                    "non invia alcun byte sulla rete; in modalità hybrid "
                    "non vengono trasmessi cookie, credenziali o body "
                    "di form, ma solo fingerprint del browser.")
    add_bullet(doc, "Osservabilità: ogni evento viene loggato in formato "
                    "strutturato e fan-out via WebSocket alla dashboard.")

    add_heading(doc, "2.4 Vincoli e ipotesi di lavoro", level=2)

    add_par(doc,
        "La progettazione del sistema si è sviluppata sotto "
        "alcuni vincoli espliciti, che è opportuno enunciare "
        "per delimitare il perimetro del lavoro.")

    add_par(doc,
        "Sul versante della modellazione della minaccia, si è "
        "assunto che l'attaccante agisca a livello applicativo "
        "via web, eventualmente attraverso tool industriali "
        "(Evilginx, Modlishka, EvilProxy o framework BitM+ "
        "ad-hoc), e che non abbia preventivamente compromesso il "
        "browser della vittima con un trojan locale. Quest'ultimo "
        "scenario rientra infatti nel modello MitB ed è "
        "indirizzato da soluzioni endpoint-protection diverse.")

    add_par(doc,
        "Sul versante deployment, si è ipotizzato che il "
        "servizio venga eseguito dietro un reverse-proxy che "
        "termina TLS e inoltra le richieste a uvicorn. La "
        "gestione di X-Forwarded-For, configurabile via la "
        "variabile TRUSTED_PROXIES, è coerente con questo "
        "schema.")

    add_par(doc,
        "Sul versante economico, per i deployment che "
        "utilizzano il backend Anthropic si assume "
        "l'accettabilità di un costo per richiesta dell'ordine "
        "del decimo di centesimo (target raggiunto utilizzando "
        "Claude Haiku come modello di default e un prompt "
        "compatto). Per deployment che richiedono costi "
        "ricorrenti nulli o privacy stringente, è disponibile "
        "il backend Ollama con llama3.1.")

    add_page_break(doc)


def capitolo3_progettazione(doc):
    add_capitolo(doc, "IV", "Progettazione dell'architettura")

    add_heading(doc, "3.1 Visione d'insieme della pipeline", level=2)

    add_par(doc,
        "L'architettura del sistema BitM-LLM è organizzata "
        "intorno a una singola pipeline di richiesta che "
        "attraversa nove stadi distinti, ciascuno responsabile "
        "di una specifica funzione. La separazione netta delle "
        "responsabilità è una scelta progettuale dettata sia da "
        "considerazioni di manutenibilità (ogni stadio è un "
        "modulo Python autonomo, testabile in isolamento) sia "
        "da considerazioni di performance (gli stadi più "
        "costosi possono essere short-circuitati quando lo "
        "stadio precedente ha già prodotto un verdetto "
        "definitivo).")

    add_par(doc, "La pipeline si compone come segue:")

    add_code(doc,
        "HTTP POST /api/bitm/collect\n"
        " │\n"
        " ├─ 1. GeoIP middleware       → arricchisce con country/ASN/ISP\n"
        " ├─ 2. rate_check             → sliding window; 429 se superato\n"
        " ├─ 3. is_blocked             → controlla il set IP banditi\n"
        " ├─ 4. session load/merge     → recupera sessione da Redis o memory\n"
        " ├─ 5. extract_features       → calcola pre_risk_score + signals\n"
        " ├─ 6. _fast_rules            → regole deterministiche (0 ms)\n"
        " ├─ 7. score_session (LLM)    → scoring Anthropic / Ollama / stub\n"
        " ├─ 8. analyze_trajectory     → secondo layer LLM su pages+timings\n"
        " ├─ 9. decide (policy)        → action finale con boost contestuali\n"
        " ├─ persistenza sessione      → set su Redis o dict in-memory\n"
        " ├─ log_event                 → JSONL append + stdout colorato\n"
        " ├─ broadcaster.publish       → fan-out WebSocket /ws/events\n"
        " └─ notify_block              → webhook fire-and-forget se BLOCK\n"
    )

    add_figure(doc, fig_pipeline_diagram(),
               "Figura 3.1 — Pipeline a nove stadi della richiesta "
               "/api/bitm/collect, con i blocchi colorati per "
               "responsabilità (rete, fingerprint, scoring, decisione).")

    add_par(doc,
        "Il punto di ingresso unico per il client è l'endpoint "
        "POST /api/bitm/collect, che riceve un payload JSON "
        "contenente il fingerprint del browser (User-Agent, "
        "plugin, WebGL renderer, canvas hash, lingue, timezone, "
        "screen, page, timing, sessionId opzionale) e restituisce "
        "una risposta JSON contenente l'azione decisa, lo score, "
        "il verdetto, la confidenza, la lista degli indicatori, "
        "la motivazione, il contesto e la latenza misurata.")

    add_heading(doc, "3.2 Modello di scoring a due stadi", level=2)

    add_par(doc,
        "Una scelta architetturale centrale, e probabilmente "
        "quella che più caratterizza il contributo di questo "
        "lavoro rispetto ai sistemi commerciali esistenti, è "
        "il modello di scoring a due stadi. L'idea è di "
        "separare nettamente i segnali deterministici "
        "(verificabili con regole esatte) dai segnali "
        "probabilistici (richiedenti una valutazione "
        "qualitativa) e di garantire che i primi non possano "
        "mai essere \"argomentati via\" dai secondi.")

    add_par(doc,
        "Concretamente, il modulo extractor.py calcola "
        "innanzitutto un valore numerico denominato "
        "pre_risk_score, ottenuto come somma pesata di tutti "
        "i segnali confermati. I pesi sono assegnati in modo "
        "che un singolo segnale critico (ad esempio "
        "headlesschrome_ua = 0.50, webdriver_true = 0.45, "
        "novnc_client_marker = 0.80) sia sufficiente, da solo, "
        "a generare uno score abbastanza elevato da forzare "
        "un'azione di BLOCK indipendentemente dal contributo "
        "del LLM.")

    add_par(doc,
        "Il valore di pre_risk_score viene poi passato come "
        "input al LLM (che può quindi tenerne conto nella "
        "propria valutazione qualitativa) e successivamente "
        "trattato da policy.py come floor: lo score finale "
        "utilizzato per la decisione è il massimo fra "
        "pre_risk_score e il valore restituito dal LLM. Il "
        "rationale di questa scelta è duplice: da un lato "
        "previene il caso patologico in cui l'LLM, per un "
        "errore di prompt-injection o per un'allucinazione, "
        "decida di scagionare un browser che presenta "
        "indicatori inequivocabili di automazione; dall'altro "
        "permette al LLM di alzare lo score quando rileva "
        "combinazioni non-lineari di segnali che il pre_score "
        "deterministico, per costruzione lineare, non potrebbe "
        "cogliere.")

    add_figure(doc, fig_two_stage_scoring(),
               "Figura 3.2 — Modello di scoring a due stadi: il pre_risk_score "
               "deterministico funge da floor sullo score restituito dal LLM, "
               "impedendo che un'allucinazione del modello scagioni segnali "
               "certi.")

    add_figure(doc, fig_signal_weights(),
               "Figura 3.3 — Pesi dei segnali deterministici utilizzati da "
               "extractor._pre_score. I marker BitM/BitM+ ad alta diagnostica "
               "(0.70–0.80) sono dimensionati per superare da soli la soglia "
               "di BLOCK del contesto più permissivo.")

    add_par(doc,
        "Una tipologia particolare di segnali, denominata "
        "CRITICAL_BLOCK, gode di trattamento speciale: la "
        "presenza anche di un solo indicatore appartenente a "
        "questo insieme forza immediatamente l'azione BLOCK, "
        "senza nemmeno consultare il LLM. Fanno parte di "
        "questo insieme sia segnali di automazione conclamata "
        "(headlesschrome_ua, webdriver_true, "
        "no_plugins_no_webgl, extreme_latency, tor_exit_node) "
        "sia i marker BitM/BitM+ ad alta diagnostica "
        "(novnc_client_marker, guacamole_client_marker, "
        "bitm_framework_ua, bitm_backend_port, "
        "xss_reflected_param, webauthn_api_override, "
        "bitm_websocket_transport).")

    add_heading(doc, "3.3 Soglie contestuali e politica decisionale", level=2)

    add_par(doc,
        "Lo score finale, una volta calcolato, viene "
        "confrontato con due soglie — challenge e block — il "
        "cui valore dipende dal contesto della pagina richiesta. "
        "La motivazione di questa differenziazione è che lo "
        "stesso comportamento può essere innocuo o sospetto a "
        "seconda della pagina: un browser senza WebGL su una "
        "pagina di documentazione ha probabilità trascurabile "
        "di essere malevolo, lo stesso browser su una pagina "
        "di pagamento merita un'analisi più stretta.")

    add_par(doc,
        "Le soglie effettivamente adottate, frutto di "
        "calibrazione empirica sulla suite di test, sono:")

    add_code(doc,
        "THRESHOLDS = {\n"
        '    "default": (0.40, 0.75),\n'
        '    "login":   (0.28, 0.62),\n'
        '    "payment": (0.20, 0.55),\n'
        '    "admin":   (0.22, 0.60),\n'
        '    "static":  (0.70, 0.92),\n'
        "}\n"
    )

    add_par(doc,
        "Il primo valore di ogni coppia rappresenta la soglia "
        "di CHALLENGE, il secondo la soglia di BLOCK. "
        "Il contesto static, riservato a richieste per asset "
        "statici (CSS, immagini, font), è stato volutamente "
        "calibrato verso l'alto per evitare di disturbare il "
        "rendering delle pagine: nessun attaccante BitM ha "
        "interesse a invocare /static/logo.png in modo "
        "anomalo.")

    add_figure(doc, fig_thresholds(),
               "Figura 3.4 — Soglie contestuali: bande ALLOW/CHALLENGE/BLOCK "
               "per ciascuna categoria di pagina. Le pagine sensibili "
               "(payment, admin, login) hanno soglie più strette del default; "
               "le pagine static sono volutamente permissive.")

    add_par(doc,
        "In aggiunta alle soglie, in contesti sensibili "
        "(login, payment, admin) viene applicato un boost "
        "moltiplicativo allo score, ottenuto sommando pesi "
        "individuali per ciascun segnale debole presente "
        "(VPN +0.16, timezone_anomaly +0.12, swiftshader_webgl "
        "+0.10, e così via). La somma totale del boost è "
        "limitata al valore MAX_BOOST = 0.25, una scelta "
        "esplicitamente conservativa che impedisce a una "
        "qualsiasi catena di segnali deboli di portare un "
        "browser legittimo sopra la soglia di BLOCK.")

    add_heading(doc, "3.4 Layer Trajectory: l'analisi di sessione", level=2)

    add_par(doc,
        "Il fingerprint di un singolo browser, per quanto "
        "ricco, è un'istantanea: cattura uno stato in un "
        "momento, ma non racconta la storia. Eppure molti "
        "attacchi BitM si rivelano proprio nelle storie: una "
        "sessione che dopo essersi loggata cambia password "
        "in tre secondi, un IP che accede direttamente a "
        "/admin senza essere mai passato da /login, un "
        "browser che apre venti pagine in un secondo. "
        "L'introduzione del layer Trajectory, avvenuta nella "
        "versione 7.4 del sistema, risponde a questo gap.")

    add_par(doc,
        "Il layer è realizzato come una seconda chiamata LLM "
        "indipendente da quella di scoring del fingerprint. "
        "L'input passato al modello è una rappresentazione "
        "strutturata della sessione corrente: la lista delle "
        "ultime dieci pagine visitate, i tempi di richiesta, "
        "il timestamp di prima visione, oltre al pre_risk_score "
        "e ai confirmed_signals del fingerprint per dare al "
        "modello il contesto completo. L'output atteso è un "
        "JSON con quattro campi: trajectory_score (0-1), "
        "pattern (snake_case fra panic_password_change, "
        "direct_admin_access, rapid_navigation, normal_flow), "
        "explanation_user (italiano, ≤160 caratteri, rivolta "
        "all'utente finale), explanation_admin (≤200 "
        "caratteri, tecnica, rivolta all'operatore SOC).")

    add_par(doc,
        "Il trajectory_score così ottenuto contribuisce alla "
        "decisione finale come boost addizionale, capato a "
        "TRAJ_BOOST_CAP = 0.25. È importante notare che il cap "
        "è separato da quello dei segnali deboli: di proposito, "
        "perché un pattern di traiettoria forte (come il "
        "panic password change) deve essere in grado di "
        "spingere uno score borderline sopra la soglia di "
        "challenge anche se i segnali fingerprint sono puliti. "
        "Resta tuttavia inferiore alla soglia di BLOCK del "
        "contesto admin (0.60), il che significa che la sola "
        "traiettoria, senza alcun fingerprint sospetto, non "
        "può forzare un blocco — una scelta deliberatamente "
        "conservativa per limitare i falsi positivi.")

    add_figure(doc, fig_trajectory_patterns(),
               "Figura 3.5 — Pattern di traiettoria riconosciuti dal layer "
               "Trajectory e relativo trajectory_score. La soglia di sospetto "
               "(0.21) attiva il boost contestuale; oltre 0.51 lo score è "
               "considerato indicatore forte di compromissione.")

    add_par(doc,
        "Per ridurre il costo del layer, è stato implementato "
        "uno short-circuit deterministico: se la sessione "
        "contiene meno di due pagine visitate, il pattern "
        "viene immediatamente impostato a insufficient_history "
        "e nessuna chiamata LLM viene effettuata. "
        "Analogamente, se nessuna delle pagine visitate "
        "appartiene a categorie sensibili (login, admin, "
        "change-password), la chiamata viene saltata e "
        "il pattern impostato a normal_flow. Questo "
        "ottimizzazione, introdotta in v7.4.2, elimina circa "
        "1 secondo di round-trip LLM su ogni sessione "
        "\"noiosa\".")

    add_heading(doc, "3.5 Persistenza, sessioni e rate-limiting", level=2)

    add_par(doc,
        "La gestione dello stato è affidata a un singolo "
        "componente, denominato SessionStore, che astrae le "
        "tre tipologie di stato persistente del sistema: le "
        "sessioni utente (con TTL configurabile, default 1 "
        "ora), il set degli IP banditi (senza TTL), le "
        "finestre di rate-limit (sliding window di 60 "
        "secondi).")

    add_par(doc,
        "Il SessionStore implementa internamente due backend: "
        "uno basato su Redis e uno basato su strutture dati "
        "Python in-memory. La scelta del backend è "
        "trasparente al chiamante: la classe espone un'API "
        "asincrona omogenea che internamente, in caso di "
        "fallimento di Redis, ricade automaticamente sul "
        "backend in-memory. Questa scelta è motivata dal "
        "requisito non funzionale di disponibilità: un "
        "outage di Redis non deve interrompere il servizio, "
        "ma al massimo limitarne la condivisione di stato "
        "fra più worker.")

    add_par(doc,
        "Il rate-limiting è realizzato con una struttura a "
        "sliding window: per ogni IP viene mantenuto un "
        "sorted-set Redis (o un deque Python) contenente i "
        "timestamp delle richieste accettate negli ultimi 60 "
        "secondi. Una nuova richiesta è accettata se il "
        "conteggio è inferiore a RATE_LIMIT (default 30), e "
        "in tal caso il timestamp viene aggiunto al set. La "
        "scelta di non contare le richieste rifiutate, "
        "introdotta in v7.4.1, evita il caso patologico in "
        "cui un client aggressivo potrebbe restare bloccato "
        "all'infinito anche dopo aver ridotto il proprio "
        "rate.")

    add_par(doc,
        "L'escalation di blocco è una funzione di lungo "
        "periodo: se la stessa sessione raccoglie tre BLOCK "
        "consecutivi, l'IP corrispondente viene aggiunto al "
        "set degli IP banditi e tutte le sue richieste "
        "successive ricevono immediatamente un BLOCK con "
        "indicatore ip_previously_blocked. La rimozione "
        "dell'IP dal set richiede un'azione amministrativa "
        "esplicita via DELETE /api/bitm/sessions.")

    add_heading(doc, "3.6 Estensione browser BitM Shield", level=2)

    add_par(doc,
        "A complemento del backend server-side, il progetto "
        "include un'estensione browser MV3 (Manifest V3) "
        "denominata BitM Shield, che porta la stessa logica "
        "di rilevamento direttamente nel browser dell'utente "
        "finale. La motivazione è di natura distributiva: "
        "non tutti i siti integreranno il backend, ma "
        "l'utente che desidera proteggersi può installare "
        "l'estensione e ottenere un livello di difesa "
        "indipendente dalla volontà del sito.")

    add_par(doc,
        "BitM Shield implementa tre modalità operative, "
        "selezionabili dal popup: off (estensione "
        "disattivata), local (rilevamento puramente locale, "
        "nessuna rete), hybrid (POST del fingerprint al "
        "backend per ottenere spiegazioni LLM e analisi di "
        "traiettoria). La modalità di default è local, "
        "scelta che privilegia la privacy: in local mode "
        "nessun byte lascia il browser dell'utente.")

    add_par(doc,
        "L'estensione realizza il rilevamento attraverso "
        "due content script. Il primo, page-hook.js, viene "
        "iniettato a document_start nel main world e si "
        "occupa di intercettare la creazione di WebSocket, "
        "le interrogazioni a navigator.credentials.get(), "
        "il rendering del canvas e la query del WebGL "
        "renderer. Il secondo, content-script.js, è "
        "iniettato nell'isolated world e implementa la "
        "pipeline detect → optional probe → banner.")

    add_par(doc,
        "Quando viene rilevato un attacco, l'estensione "
        "presenta nel DOM della pagina un banner shadow-DOM "
        "che spiega all'utente la natura del problema in "
        "italiano (o in inglese, in base al locale del "
        "browser) e blocca preventivamente l'invio di form "
        "che contengano campi password. Il badge dell'icona "
        "dell'estensione cambia colore in funzione del "
        "verdetto: grigio per allow, arancione con punto "
        "esclamativo per challenge, rosso con croce per "
        "block.")

    add_par(doc,
        "Una funzionalità di hardening aggiuntiva, "
        "selezionabile via popup, consiste nell'installazione "
        "di regole declarativeNetRequest che bloccano "
        "richieste di rete verso i domini di tunneling più "
        "comunemente sfruttati dal BitM+ (ngrok, "
        "trycloudflare, localtunnel, serveo) ancor prima "
        "che la pagina venga renderizzata. Questa misura, "
        "essendo applicata a livello di rete, è "
        "particolarmente efficace contro varianti che "
        "tentano di esfiltrare dati attraverso WebSocket "
        "diretti a tunnel pubblici.")

    add_page_break(doc)


def capitolo4_implementazione(doc):
    add_capitolo(doc, "V", "Implementazione")

    add_heading(doc, "4.1 Stack tecnologico e organizzazione del codice",
                level=2)

    add_par(doc,
        "Il sistema è stato implementato in Python 3.11+ "
        "utilizzando il framework FastAPI come server HTTP "
        "asincrono. La scelta di FastAPI è motivata dal suo "
        "supporto nativo per asyncio (necessario per gestire "
        "in modo non-bloccante le chiamate al LLM e a Redis), "
        "dalla generazione automatica della documentazione "
        "OpenAPI, e dall'integrazione fluida con le librerie "
        "di logging strutturato e validazione Pydantic.")

    add_par(doc,
        "Le librerie di terze parti rilevanti sono: anthropic "
        "(client ufficiale Anthropic), httpx (client HTTP "
        "asincrono utilizzato per Ollama e per i webhook), "
        "redis (client asincrono per il SessionStore), "
        "geoip2 (lettura dei database MaxMind GeoLite2). "
        "Tutte le dipendenze sono dichiarate in requirements.txt.")

    add_par(doc,
        "L'organizzazione del codice nel pacchetto bitm-plugin/app "
        "rispecchia direttamente la pipeline architetturale:")

    add_code(doc,
        "bitm-plugin/\n"
        "├── app/\n"
        "│   ├── main.py          # FastAPI entry point + middleware GeoIP\n"
        "│   ├── config.py        # Variabili d'ambiente, validazione\n"
        "│   ├── extractor.py     # pre_risk_score, confirmed_signals, BitM\n"
        "│   ├── scorer.py        # Backend LLM: Anthropic / Ollama / stub\n"
        "│   ├── policy.py        # Soglie, boost, fast-track, decide()\n"
        "│   ├── geoip.py         # MaxMind resolver + VPN/Tor detection\n"
        "│   ├── redis_client.py  # SessionStore: Redis + fallback in-memory\n"
        "│   ├── broadcaster.py   # Pub/sub WebSocket per dashboard\n"
        "│   ├── notifier.py      # Webhook push asincrono\n"
        "│   ├── logger.py        # Logging stdout + JSONL\n"
        "│   └── static/          # collector.js, dashboard.html, test_page.html\n"
        "├── tests/\n"
        "│   └── run_tests.py     # Suite 49 casi (29 funzionali + 20 sistema)\n"
        "├── training/\n"
        "│   ├── build_dataset.py # bitm_events.jsonl → dataset ChatML\n"
        "│   └── train_lora.py    # Fine-tuning LoRA Llama 3.1\n"
        "├── run.py               # Entry point uvicorn\n"
        "├── diagnose.py          # Diagnostica end-to-end\n"
        "└── requirements.txt\n"
    )

    add_heading(doc, "4.2 Il modulo extractor.py", level=2)

    add_par(doc,
        "Il modulo extractor.py è responsabile della "
        "trasformazione del payload JSON ricevuto dal client "
        "in una struttura dati arricchita, che costituisce "
        "l'input degli stadi successivi della pipeline. La "
        "funzione pubblica principale è extract_features, che "
        "accetta come parametri il dizionario raw del payload, "
        "l'IP risolto, lo stato di sessione corrente e i "
        "metadati GeoIP, e restituisce un dizionario con "
        "circa trenta campi.")

    add_par(doc,
        "All'interno della funzione vengono calcolati tre "
        "blocchi di feature: identità del browser e del "
        "sistema operativo (parsing dello User-Agent), "
        "fingerprint propriamente detto (plugin, WebGL, "
        "canvas hash, lingue, screen, timezone), metriche "
        "di sessione (medie, massimi e deviazione standard "
        "dei tempi di richiesta).")

    add_par(doc,
        "Le funzioni interne più significative sono tre. "
        "La prima, _detect_headless, restituisce la lista "
        "dei segnali compatibili con un browser automatizzato "
        "(headlesschrome_ua, webdriver_true, zero_plugins, "
        "no_webgl_renderer, swiftshader_webgl, empty_canvas, "
        "no_languages, suspicious_resolution, low_color_depth, "
        "no_timezone). La seconda, _detect_bitm, restituisce "
        "la lista dei segnali specifici degli stack BitM e "
        "BitM+ documentati nei papers di Tommasi, Tzschoppe e "
        "Catalano. La terza, _pre_score, calcola il valore "
        "numerico pre_risk_score sommando i pesi dei segnali "
        "rilevati dalle due funzioni precedenti.")

    add_par(doc,
        "La sezione di rilevamento BitM merita un commento "
        "specifico perché incorpora la maggior parte del "
        "valore aggiunto rispetto a un classico fingerprinter. "
        "Le firme implementate sono:")

    add_bullet(doc, "Tunnel HTTPS verso localhost: regex su pageUrl "
                    "e referrer per riconoscere domini "
                    "ngrok-free.app, trycloudflare.com, loca.lt, "
                    "localtunnel.me, serveo.net.")
    add_bullet(doc, "Marker nel document.title: \"noVNC\" o "
                    "\"Websockify\" generano novnc_client_marker; "
                    "\"Guacamole\" genera guacamole_client_marker. "
                    "I check sono soppressi se il titolo "
                    "appartiene a una pagina di motore di ricerca, "
                    "per evitare il falso positivo \"noVNC - "
                    "Ricerca Google\".")
    add_bullet(doc, "Porte di backend BitM+: regex su pageUrl per "
                    "riconoscere :3081 (MalSrv Express), :6080 "
                    "(noVNC), :4822 (Guacamole Tomcat), :5900 "
                    "(VNC nativo).")
    add_bullet(doc, "Payload XSS riflesso: regex per pattern "
                    "tipici di evilGet/loadFromAttacker — "
                    "<script>, onerror=, javascript:, "
                    "document.createElement, eval, "
                    "fromCharCode.")
    add_bullet(doc, "Override WebAuthn: il client può fornire "
                    "credentialsGetNative=false se la verifica "
                    "navigator.credentials.get.toString() === "
                    "\"function get() { [native code] }\" fallisce.")
    add_bullet(doc, "WebSocket transport sospetto: gli endpoint "
                    "WebSocket aperti dalla pagina vengono "
                    "ispezionati per rilevare path /websockify, "
                    "/vnc, /guacamole o tunnel.")
    add_bullet(doc, "User-Agent compromessi: presenza di stringhe "
                    "noVNC, Websockify, Guacamole, TigerVNC nello "
                    "User-Agent del browser viewer interno (PoC "
                    "non-stealth).")
    add_bullet(doc, "Iframe overlay: rilevamento di pagine con "
                    "5 o più iframe, pattern tipico dello "
                    "Scenario 3 di BitM+ in cui un iframe "
                    "full-screen copre il viewport.")

    add_heading(doc, "4.3 Il modulo scorer.py: integrazione di Anthropic e Ollama",
                level=2)

    add_par(doc,
        "Il modulo scorer.py è il componente più articolato "
        "del sistema e implementa l'interfaccia unificata "
        "verso i tre backend LLM supportati: Anthropic Claude, "
        "Ollama locale e stub deterministico. La funzione "
        "pubblica principale è score_session, che riceve in "
        "input le feature calcolate da extractor e restituisce "
        "un dizionario con quattro campi: risk_score, verdict, "
        "confidence, indicators, explanation.")

    add_par(doc,
        "Tutti e tre i backend condividono lo stesso prompt "
        "di sistema, una versione compatta — circa 640 "
        "caratteri, 40% inferiore alla v6 — che istruisce il "
        "modello a rispondere esclusivamente con un oggetto "
        "JSON conforme a uno schema preciso, e definisce le "
        "soglie di mappatura risk_score → verdict (0.00-0.30 "
        "LEGITIMATE, 0.31-0.64 SUSPICIOUS, 0.65-1.00 ATTACK). "
        "Lo stesso prompt esplicita il vincolo di floor sul "
        "pre_risk_score, che il modello deve rispettare.")

    add_par(doc,
        "Il prompt utente, generato dalla funzione _build_prompt, "
        "è un blocco testuale strutturato in cinque sezioni: "
        "punteggio deterministico pre-calcolato, dettagli del "
        "browser, rete e timing, comportamento, istruzioni "
        "finali. La struttura mantiene volutamente i nomi "
        "delle feature in inglese (matching del SYSTEM_PROMPT) "
        "ma le etichette di sezione in italiano per coerenza "
        "con il resto del sistema.")

    add_par(doc,
        "Il backend Anthropic è implementato dalla funzione "
        "_score_anthropic. Il client utilizzato è "
        "anthropic.AsyncAnthropic. La selezione del modello "
        "avviene dinamicamente al primo utilizzo: la funzione "
        "_anthropic_pick_model itera sulla lista "
        "ANTHROPIC_MODELS (Claude Haiku 4.5, Claude 3.5 Haiku, "
        "Claude Sonnet 4.6, Claude 3.5 Sonnet, Claude 3 Haiku) "
        "e seleziona il primo che risponde a un ping di test. "
        "Questa logica garantisce che il sistema continui a "
        "funzionare anche se uno dei modelli viene deprecato. "
        "La gestione degli errori prevede tre tentativi con "
        "backoff esponenziale per gli errori 5xx; le risposte "
        "non-JSON e gli errori di rete vengono riportati come "
        "errori \"morbidi\" con risk_score 0.5 e indicator "
        "specifico (llm_parse_error, api_error).")

    add_par(doc,
        "Il backend Ollama è implementato dalla funzione "
        "_score_ollama. Si appoggia all'API REST nativa di "
        "Ollama (/api/chat) tramite httpx. Una specificità "
        "rilevante è l'uso del parametro \"format\": \"json\" "
        "supportato da llama3.1, che forza il modello a "
        "produrre esclusivamente output JSON valido senza "
        "alcun preambolo. Il parser è comunque difensivo, "
        "in grado di rimuovere blocchi markdown ```json...``` "
        "e di estrarre il JSON outermost in caso di testo "
        "introduttivo.")

    add_par(doc,
        "Il backend stub, infine, è una pura derivazione "
        "deterministica dei segnali calcolati da extractor: "
        "score = pre_risk_score, verdict mappato sulle "
        "soglie LEGITIMATE/SUSPICIOUS/ATTACK, indicators "
        "ottenuti unendo confirmed_signals e headless_signals. "
        "Questo backend è il default in assenza di "
        "configurazione e ha tre vantaggi pratici: nessuna "
        "API key richiesta per il primo avvio, nessuna "
        "dipendenza esterna, riproducibilità perfetta per "
        "i test E2E.")

    add_par(doc,
        "Indipendentemente dal backend, ogni risposta passa "
        "attraverso la funzione _validate_result che applica "
        "due tipi di normalizzazione: validazione dei tipi "
        "(risk_score clipping in [0, 1], verdict negli enum "
        "validi, confidence negli enum validi, indicators "
        "lista di stringhe) e coerenza fra verdict e score "
        "(un LEGITIMATE con score ≥ 0.65 viene promosso ad "
        "ATTACK, un ATTACK con score ≤ 0.30 viene declassato "
        "a SUSPICIOUS).")

    add_par(doc,
        "Una cache TTL condivisa fra i backend, keyed su "
        "(canvas_hash, user_agent[:60]), evita di ripetere "
        "la chiamata LLM per fingerprint identici entro la "
        "finestra di CACHE_TTL secondi (default 300). Solo "
        "le risposte di successo vengono cacheate; gli "
        "errori vengono propagati al chiamante con il flag "
        "_from_cache assente.")

    add_heading(doc, "4.4 Il modulo policy.py", level=2)

    add_par(doc,
        "Il modulo policy.py traduce lo score numerico "
        "prodotto dallo scorer in una decisione operativa. "
        "Le funzioni pubbliche sono due: detect_page_context, "
        "che mappa l'URL della richiesta a una delle cinque "
        "categorie di contesto, e decide, che restituisce "
        "una tupla (Action, motivazione).")

    add_par(doc,
        "La funzione decide implementa una sequenza di "
        "controlli in cinque passi, ciascuno responsabile di "
        "una specifica trasformazione dello score. Il primo "
        "passo verifica la presenza di segnali critici "
        "nell'unione fra gli indicatori del LLM e i "
        "confirmed_signals di extractor: se l'intersezione "
        "con CRITICAL_BLOCK non è vuota, il verdetto è "
        "immediatamente BLOCK con score forzato a 0.97. "
        "Il secondo passo applica il floor del pre_risk_score: "
        "se pre_score > score, allora score = pre_score. "
        "Il terzo passo applica il boost contestuale dei "
        "segnali deboli, capato a MAX_BOOST = 0.25. Il "
        "quarto passo applica il boost trajectory, capato "
        "indipendentemente a TRAJ_BOOST_CAP = 0.25. Il "
        "quinto passo confronta lo score amplificato con le "
        "soglie del contesto e restituisce ALLOW, CHALLENGE "
        "o BLOCK.")

    add_par(doc,
        "Una sottigliezza implementativa che merita di "
        "essere menzionata riguarda il parsing dello score: "
        "in Python il valore 0.0 è falsy, e la scrittura "
        "naturale score = score_result.get(\"risk_score\") "
        "or 0.5 produrrebbe 0.5 ogni volta che lo score "
        "vero è 0.0, trasformando silenziosamente i casi "
        "legittimi in casi sospetti. Il codice usa quindi "
        "un controllo esplicito su None, una pratica che "
        "abbiamo formalizzato come invariante di "
        "manutenzione del modulo.")

    add_heading(doc, "4.5 Il SessionStore: Redis con fallback in-memory",
                level=2)

    add_par(doc,
        "La classe SessionStore, definita in redis_client.py, "
        "fornisce un'unica interfaccia asincrona per la "
        "persistenza di sessioni, IP banditi e finestre di "
        "rate-limit. Il pattern adottato è quello del "
        "graceful degradation: ogni metodo che effettua una "
        "chiamata Redis è racchiuso in un try/except che, "
        "in caso di errore, marca la connessione come non "
        "disponibile e ricade automaticamente sull'equivalente "
        "in-memory.")

    add_par(doc,
        "Le chiavi Redis sono prefissate dalla variabile "
        "REDIS_KEY_PREFIX (default \"bitm:\") per evitare "
        "collisioni con altri servizi che condividano lo "
        "stesso database. Lo schema delle chiavi è:")

    add_bullet(doc, "{prefix}session:{sid} → hash JSON con TTL "
                    "REDIS_SESSION_TTL (default 3600 s).")
    add_bullet(doc, "{prefix}blocked → set di IP banditi (senza TTL).")
    add_bullet(doc, "{prefix}rate:{ip} → sorted-set con "
                    "(timestamp → timestamp) e TTL pari alla finestra.")

    add_par(doc,
        "Il rate-limiting su Redis è realizzato con una "
        "pipeline bifase. La prima fase elimina i timestamp "
        "scaduti e conta quelli rimanenti; se il conteggio "
        "supera la soglia, la richiesta è rifiutata. La "
        "seconda fase, eseguita solo in caso di "
        "accettazione, aggiunge il nuovo timestamp e "
        "rinfresca il TTL. La separazione delle due fasi "
        "previene il problema delle richieste rifiutate "
        "che gonfiano la finestra: anomalia osservata e "
        "corretta in v7.4.1.")

    add_heading(doc, "4.6 Real-time dashboard e webhook push", level=2)

    add_par(doc,
        "Il modulo broadcaster.py implementa un semplice "
        "pub/sub in-process per la dashboard real-time. La "
        "classe EventBroadcaster mantiene un set di client "
        "WebSocket connessi a /ws/events e un ring buffer "
        "circolare di 500 eventi recenti. Ogni richiesta a "
        "/api/bitm/collect produce, alla fine della pipeline, "
        "una chiamata a broadcaster.publish che inoltra "
        "l'evento serializzato in JSON a tutti i client "
        "connessi. I nuovi client, al momento della "
        "connessione, ricevono il backlog corrente in modo "
        "che la dashboard non parta vuota.")

    add_par(doc,
        "La dashboard, accessibile via /dashboard "
        "(autenticata con ADMIN_TOKEN), è un'applicazione "
        "HTML/JavaScript single-file che si connette al "
        "WebSocket, mantiene una propria buffer di eventi "
        "in memoria, e renderizza in tempo reale due grafici "
        "Chart.js (distribuzione delle azioni nel tempo, top "
        "indicators) oltre a una tabella scorrevole degli "
        "eventi più recenti. Un pulsante consente l'export "
        "CSV degli eventi in buffer.")

    add_par(doc,
        "Una limitazione architetturale del broadcaster, "
        "dichiarata in modo esplicito nel codice, è che il "
        "pub/sub è in-process e quindi single-worker: in un "
        "deployment con uvicorn --workers > 1 ciascun worker "
        "manterrebbe il proprio set di client e ciascun "
        "evento sarebbe visibile solo dai client connessi al "
        "worker che ha servito la richiesta. Per scenari "
        "multi-worker è necessario promuovere il trasporto a "
        "Redis pub/sub, una modifica che non è stata "
        "implementata in questo lavoro perché esula dalle "
        "esigenze del progetto.")

    add_par(doc,
        "Il modulo notifier.py implementa il push asincrono "
        "verso webhook esterni in caso di evento BLOCK. La "
        "configurazione è leggibile da variabili d'ambiente "
        "(WEBHOOK_URL, WEBHOOK_TYPE, WEBHOOK_TIMEOUT, "
        "WEBHOOK_RETRIES) o da un file JSON dedicato "
        "(WEBHOOK_CONFIG_FILE). Sono supportati tre formati "
        "di payload: Slack (Blocks API), Microsoft Teams "
        "(Adaptive Cards v1.4) e SIEM generico (JSON "
        "strutturato standard).")

    add_par(doc,
        "L'invio è implementato come fire-and-forget tramite "
        "asyncio.create_task: la risposta al client originale "
        "non attende il completamento del webhook. Un retry "
        "esponenziale con backoff fino a 30 secondi gestisce "
        "gli errori transitori; gli errori 4xx (client error) "
        "non vengono ritentati. I task pendenti sono "
        "mantenuti in un set di strong reference per evitare "
        "che il GC li raccolga prima del completamento.")

    add_heading(doc, "4.7 Hardening di sicurezza e considerazioni di deploy",
                level=2)

    add_par(doc,
        "Il sistema integra alcune misure di hardening "
        "specificamente progettate per il deployment in "
        "produzione. La prima è il trattamento sicuro "
        "dell'header X-Forwarded-For: per evitare che un "
        "attaccante possa ruotare il proprio IP "
        "manipolando questo header, il middleware GeoIP lo "
        "considera valido solo se il peer diretto della "
        "connessione TCP appartiene alla lista "
        "TRUSTED_PROXIES. In caso contrario viene utilizzato "
        "request.client.host. La lista è configurabile come "
        "CSV di IP o CIDR via .env.")

    add_par(doc,
        "La seconda misura è la protezione degli endpoint "
        "amministrativi (GET/DELETE /api/bitm/sessions, "
        "/dashboard, /ws/events) tramite un token "
        "configurabile via ADMIN_TOKEN. La verifica del "
        "token è implementata con hmac.compare_digest per "
        "evitare timing attack. Il token può essere fornito "
        "via header X-Admin-Token (per le chiamate JSON) o "
        "via query parameter ?token= (per la dashboard "
        "browser e per il WebSocket, dove il browser non "
        "consente di settare header custom sull'handshake).")

    add_par(doc,
        "La terza misura è l'isolamento del session ID: in "
        "assenza di un sessionId esplicito fornito dal "
        "client, il sistema genera un ID di fallback come "
        "hash SHA-1 di una concatenazione di IP, "
        "User-Agent, canvas hash e lingue. Questo schema "
        "garantisce che due client distinti dietro lo stesso "
        "NAT non condividano la sessione (e quindi non "
        "condividano l'eventuale escalation di blocco), "
        "mentre lo stesso client che effettua più richieste "
        "sequenziali viene correttamente riconosciuto.")

    add_par(doc,
        "Il deployment in produzione è supportato da un "
        "Dockerfile minimale e da un docker-compose.yml che "
        "orchestra tre servizi: bitm-plugin (l'API stessa), "
        "redis (per la persistenza condivisa) e ollama "
        "(opzionale, attivabile via profile). L'immagine "
        "Docker è disponibile su GitHub Container Registry "
        "(ghcr.io) ed è costruita automaticamente da una "
        "GitHub Action al push di ogni release.")

    add_page_break(doc)


def capitolo5_sperimentazione(doc):
    add_capitolo(doc, "VI", "Sperimentazione e validazione")

    add_heading(doc, "5.1 Metodologia di test", level=2)

    add_par(doc,
        "La validazione del sistema è stata condotta secondo "
        "una metodologia in due fasi. La prima fase, di "
        "natura unitaria, verifica il corretto funzionamento "
        "dei singoli componenti tramite una suite di 49 casi "
        "di test che esercitano l'intera pipeline da una "
        "richiesta HTTP simulata fino alla risposta JSON. La "
        "seconda fase, di natura integrativa, esercita la "
        "pipeline completa contro scenari E2E orchestrati con "
        "Playwright, in cui browser reali eseguono attacchi "
        "controllati contro un'istanza locale del sistema.")

    add_par(doc,
        "L'organizzazione della suite di test è stata "
        "progettata per coprire in modo bilanciato cinque "
        "categorie di scenari: legit (sessioni utente "
        "legittime), attack (sessioni con marker conclamati "
        "di attacco), suspicious (sessioni con segnali "
        "ambigui che richiedono challenge), edge (casi limite "
        "come payload minimi, caratteri Unicode, asset "
        "statici) e system (controlli sull'integrità "
        "dell'infrastruttura: health endpoint, rate-limit, "
        "session store, cache LLM, label alignment).")

    add_par(doc,
        "Ogni caso di test è specificato in modo dichiarativo "
        "in run_tests.py con i campi: identificatore (T01-T29 "
        "per i casi funzionali, S01-S20 per i controlli di "
        "sistema), nome descrittivo, payload JSON di input, "
        "azione attesa (allow/challenge/block, eventualmente "
        "in disgiunzione con il pipe per consentire più "
        "esiti accettabili), contesto di pagina. Il test "
        "runner esegue le richieste contro un'istanza "
        "uvicorn in ascolto su localhost:8000 e verifica "
        "che l'azione restituita sia coerente con l'azione "
        "attesa.")

    add_heading(doc, "5.2 La test-suite: 49 casi", level=2)

    add_par(doc,
        "La struttura della suite di test è riassunta nella "
        "tabella seguente. Il numero di casi per categoria "
        "riflette la priorità data alla precisione "
        "dell'azione di blocco (che ha una colonna più "
        "popolata) rispetto agli altri esiti.")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Categoria"
    hdr[1].text = "Numero casi"
    hdr[2].text = "Azione attesa principale"
    rows_data = [
        ("legit", "5 (T01–T05)", "allow"),
        ("attack", "13 (T06–T11, T21–T27)", "block"),
        ("suspicious", "6 (T12–T16, T28)", "challenge"),
        ("edge", "5 (T17–T20, T29)", "varia"),
        ("system", "20 (S01–S20)", "controlli infrastrutturali"),
    ]
    for i, (cat, n, az) in enumerate(rows_data, start=1):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = n
        table.rows[i].cells[2].text = az

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

    add_par(doc, "")

    add_figure(doc, fig_test_results(),
               "Figura 5.1 — Esito della suite di test, suddiviso per "
               "categoria. Tutti i 49 casi (29 funzionali + 20 di sistema) "
               "sono stati superati nella run di riferimento riportata in "
               "test_report.json (v7.4.2).")

    add_par(doc,
        "Alcuni casi rilevanti meritano una descrizione "
        "approfondita. T06 (HeadlessChrome / Evilginx) "
        "esercita il rilevamento del User-Agent automatizzato "
        "in un contesto di login: il payload contiene "
        "\"Mozilla/5.0 ... HeadlessChrome/120 ...\" e "
        "webdriver=true; l'esito atteso è BLOCK con "
        "indicatori headless_ua, webdriver_flag, "
        "no_plugins_no_webgl. T07 (Playwright/SwiftShader) "
        "verifica il riconoscimento di Playwright tramite la "
        "combinazione webdriver=true e WebGL=SwiftShader. "
        "T08 (Selenium / no plugin + no WebGL / desktop) "
        "esercita la regola composta no_plugins_no_webgl in "
        "un contesto di pagamento.")

    add_par(doc,
        "I casi T21–T27 sono dedicati specificamente al "
        "rilevamento BitM/BitM+ e costituiscono la principale "
        "novità della v7.2 del sistema. T21 (BitM RFB — "
        "noVNC client marker nel title) verifica che la "
        "presenza di \"noVNC\" nel document.title produca "
        "novnc_client_marker e blocco immediato; T22 esercita "
        "il marker Guacamole; T23 verifica il riconoscimento "
        "di payload XSS riflessi nell'URL; T24 verifica il "
        "rilevamento di evilGet via "
        "credentialsGetNative=false; T25 verifica il "
        "rilevamento delle porte di backend BitM+ visibili "
        "nell'URL; T26 verifica il leak dello User-Agent "
        "BitM; T27 verifica il rilevamento di endpoint "
        "WebSocket verso tunnel ngrok.")

    add_par(doc,
        "I casi suspicious (T12–T16, T28) sono progettati "
        "per esercitare il sistema in zone di valutazione "
        "borderline, dove la decisione corretta è "
        "challenge anziché block. T12 verifica che una "
        "richiesta da VPN su una pagina di login produca "
        "challenge (boost +0.16 su soglia 0.28). T13 "
        "verifica che una latenza di 380 ms su pagina di "
        "pagamento produca challenge (boost +0.12 su soglia "
        "0.20). T15 verifica il riconoscimento dell'anomalia "
        "timezone UTC + lingua italiana su pagina di "
        "amministrazione.")

    add_heading(doc, "5.3 Risultati ottenuti", level=2)

    add_par(doc,
        "L'esecuzione della suite completa, registrata in "
        "test_report.json, ha prodotto un risultato di 49 "
        "casi superati su 49, per un'accuratezza complessiva "
        "del 100%. È importante sottolineare che questo "
        "risultato è stato ottenuto con il backend Ollama "
        "llama3.1 attivo, e che la stessa suite passa "
        "egualmente con il backend Anthropic e con il backend "
        "stub deterministico — quest'ultimo, peraltro, è "
        "quello adottato in CI per eliminare la dipendenza "
        "da risorse esterne.")

    add_par(doc,
        "L'analisi delle latenze misurate è interessante "
        "perché evidenzia la natura bimodale della "
        "distribuzione. Per i casi gestiti dal fast-track "
        "deterministico (T06–T11, T17, T21–T27), la latenza "
        "totale è inferiore ai 5 ms, dominata dal solo costo "
        "di parsing JSON e dalla scrittura del log. Per i "
        "casi che richiedono la chiamata al LLM (T01–T05, "
        "T12–T16, T18–T20, T28, T29), la latenza è "
        "dell'ordine dei 1300–1800 ms con backend Ollama "
        "in locale, valore congruente con il costo medio di "
        "una chiamata all'API REST di Ollama su modello "
        "llama3.1 7B.")

    add_figure(doc, fig_score_distribution(),
               "Figura 5.2 — Distribuzione di risk_score per ciascuna azione "
               "decisa. Il box plot evidenzia come le tre azioni siano ben "
               "separate, con i casi BLOCK saturati nell'intervallo "
               "[0.97, 1.00] dai segnali critici e i casi CHALLENGE concentrati "
               "fra 0.22 e 0.50.")

    add_figure(doc, fig_latency_distribution(),
               "Figura 5.3 — Distribuzione delle latenze totali misurate "
               "sulla test-suite, in scala logaritmica. Il fast-track "
               "deterministico resta entro pochi millisecondi; la pipeline "
               "LLM completa con backend Ollama si colloca attorno ai 1.5 s, "
               "comunque sotto il target p95 di 2 s.")

    add_par(doc,
        "I casi di sistema S01–S20 verificano la corretta "
        "esposizione delle informazioni di stato (S01: "
        "/health include version, backend, store, geoip, "
        "webhook, sessioni, blocked_ips, ws_clients), la "
        "persistenza della sessione attraverso più richieste "
        "(S02), l'escalation IP dopo tre BLOCK consecutivi "
        "(S03), il funzionamento del rate-limit (S04), la "
        "gestione di IP privati senza errori (S05), "
        "l'efficacia della cache LLM (S07), il "
        "funzionamento del webhook (S08, S09), la "
        "compattezza del prompt v7 (S10, ≤ 650 caratteri), "
        "l'allineamento dei label BitM tra extractor e "
        "policy.CRITICAL_BLOCK (S13), il corretto "
        "comportamento del layer trajectory (S16–S20).")

    add_heading(doc, "5.4 Confronto fra backend LLM", level=2)

    add_par(doc,
        "Per valutare l'impatto della scelta del backend "
        "LLM sulla qualità delle decisioni, sono state "
        "eseguite tre run complete della suite, una per "
        "ciascun backend. I risultati ottenuti sono "
        "riassunti nella tabella seguente.")

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Backend"
    hdr[1].text = "Accuratezza"
    hdr[2].text = "Latenza media (ms)"
    hdr[3].text = "Costo per richiesta"
    rows_data = [
        ("Anthropic Claude Haiku 4.5", "49/49 (100%)", "~600", "~0.001 USD"),
        ("Ollama llama3.1 (locale)", "49/49 (100%)", "~1500", "0 (self-host)"),
        ("Stub deterministico", "49/49 (100%)", "~3", "0"),
    ]
    for i, (b, a, l, c) in enumerate(rows_data, start=1):
        table.rows[i].cells[0].text = b
        table.rows[i].cells[1].text = a
        table.rows[i].cells[2].text = l
        table.rows[i].cells[3].text = c
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)

    add_par(doc, "")

    add_figure(doc, fig_backend_comparison(),
               "Figura 5.4 — Confronto delle tre alternative di backend LLM "
               "su tre dimensioni: accuratezza (equivalente per costruzione "
               "dei test), latenza media (Ollama paga la propria esecuzione "
               "locale), costo per richiesta (Anthropic comporta un "
               "trascurabile costo a chiamata).")

    add_par(doc,
        "Il fatto che tutti e tre i backend ottengano la "
        "stessa accuratezza non deve sorprendere: la suite "
        "di test è dimensionata in modo tale che i casi "
        "decisivi siano coperti dal pre_risk_score "
        "deterministico (e dal floor che esso impone allo "
        "score finale) o dai segnali CRITICAL_BLOCK (che "
        "bypassano integralmente l'LLM). Il LLM interviene "
        "principalmente sui casi suspicious, dove fornisce "
        "spiegazioni in linguaggio naturale che però non "
        "modificano l'azione finale.")

    add_par(doc,
        "Una valutazione su un dataset più variegato, "
        "in cui i casi suspicious siano più rappresentati e "
        "i segnali deterministici meno discriminanti, "
        "produrrebbe verosimilmente risultati differenziati. "
        "Una stima preliminare, basata su un sottoinsieme di "
        "10 sessioni reali estratte da bitm_events.jsonl, "
        "suggerisce che Anthropic Claude Haiku produca "
        "spiegazioni più articolate e talvolta classifichi "
        "come SUSPICIOUS sessioni che Ollama llama3.1 "
        "classifica come LEGITIMATE; tuttavia il campione è "
        "troppo piccolo per trarre conclusioni statistiche "
        "robuste.")

    add_heading(doc, "5.5 Discussione dei falsi positivi e dei falsi negativi",
                level=2)

    add_par(doc,
        "Durante lo sviluppo iterativo del sistema sono "
        "stati identificati e corretti diversi tipi di "
        "falsi positivi, alcuni dei quali meritano di essere "
        "documentati perché illustrano la complessità del "
        "problema.")

    add_par(doc,
        "Un primo falso positivo riscontrato in fase di "
        "calibrazione riguardava le pagine di motori di "
        "ricerca: una ricerca per il termine \"noVNC\" "
        "produceva un document.title del tipo \"noVNC - "
        "Ricerca Google\", che attivava il segnale "
        "novnc_client_marker e generava un BLOCK "
        "indebitamente. La correzione è stata "
        "l'introduzione di un check soppressivo che "
        "disabilita il marker se il titolo contiene "
        "indicazioni di pagina di ricerca o di motore di "
        "ricerca noto (Google, Bing, DuckDuckGo, etc.).")

    add_par(doc,
        "Un secondo falso positivo riguardava browser "
        "mobili Safari su iOS, che legittimamente non hanno "
        "plugin: il segnale zero_plugins, da solo, non è "
        "discriminante. La correzione è stata l'introduzione "
        "del flag is_mobile in extract_features e la "
        "soppressione dei segnali plugin-related per "
        "dispositivi mobili.")

    add_par(doc,
        "Un terzo falso positivo, più sottile, riguardava la "
        "gestione delle latenze elevate su browser legittimi "
        "in condizioni di rete degradata. Una connessione "
        "3G lenta poteva produrre tempi di risposta nei 600 "
        "ms, attivando il segnale extreme_latency e "
        "generando un BLOCK. La correzione è stata "
        "l'introduzione di tre soglie graduate "
        "(elevated_latency >150 ms, high_latency >300 ms, "
        "extreme_latency >600 ms) con peso crescente, "
        "applicate solo come amplificatore in contesti "
        "sensibili.")

    add_par(doc,
        "Sul versante dei falsi negativi, la suite di test "
        "non ne rileva in scenari sintetici, ma una "
        "considerazione è doverosa: il sistema non può "
        "rilevare BitM perfettamente trasparenti, in cui "
        "l'attaccante abbia rimosso ogni marker BitM (titolo "
        "personalizzato, User-Agent riscritto, porte "
        "nascoste dietro un reverse-proxy) e abbia evitato "
        "tunneling pubblico in favore di un dominio "
        "registrato ad-hoc con certificato Let's Encrypt. "
        "In questo scenario worst-case, il rilevamento "
        "ricade sui soli segnali di automazione (che "
        "tuttavia restano efficaci se il browser viewer "
        "è effettivamente headless) e sui pattern di "
        "traiettoria (che intervengono dopo il primo "
        "compromesso). Una difesa più robusta richiederebbe "
        "l'integrazione con feed di reputation di domini "
        "freschi (passive DNS, certificate transparency "
        "logs) che esulano dal perimetro di questo lavoro.")

    add_page_break(doc)


def capitolo6_conclusioni(doc):
    add_capitolo(doc, "VII", "Conclusioni e sviluppi futuri")

    add_heading(doc, "6.1 Sintesi dei risultati", level=2)

    add_par(doc,
        "Il presente lavoro ha avuto come obiettivo la "
        "progettazione, l'implementazione e la validazione di "
        "un sistema di rilevamento real-time degli attacchi "
        "Browser-in-the-Middle e Browser-in-the-Middle Plus, "
        "due categorie di minaccia documentate dalla "
        "letteratura accademica come particolarmente "
        "insidiose per la loro capacità di eludere le "
        "contromisure tradizionali (URL reputation, "
        "fingerprinting statico, MFA basato su WebAuthn).")

    add_par(doc,
        "Il sistema realizzato — denominato BitM Detection "
        "Plugin, versione 7.4.2 al momento della scrittura — "
        "è un servizio FastAPI scritto in Python, organizzato "
        "in una pipeline a nove stadi che integra rilevamento "
        "deterministico, scoring statistico contestuale e "
        "valutazione tramite Large Language Model in modo "
        "complementare. La validazione su una suite di 49 "
        "casi di test (29 funzionali e 20 di sistema) ha "
        "dimostrato un'accuratezza del 100% per tutti e tre "
        "i backend LLM supportati.")

    add_par(doc,
        "I principali contributi originali del lavoro sono "
        "tre: il modello di scoring a due stadi con il "
        "pre_risk_score deterministico utilizzato come floor "
        "del verdetto LLM (che protegge il sistema da "
        "allucinazioni del modello su segnali certi); "
        "l'introduzione del layer Trajectory per il "
        "rilevamento di pattern post-compromissione che il "
        "fingerprint singolo non può catturare; "
        "l'architettura modulare a tre backend "
        "intercambiabili (Anthropic, Ollama, stub) che "
        "consente di calibrare il trade-off fra qualità del "
        "verdetto, costo di esercizio e privacy dei dati.")

    add_par(doc,
        "L'estensione browser BitM Shield, sviluppata come "
        "complemento del backend, completa il quadro "
        "fornendo una difesa lato utente che non dipende "
        "dalla cooperazione del sito visitato. Le sue "
        "modalità local e hybrid offrono un trade-off "
        "esplicito fra privacy e qualità del rilevamento, "
        "mentre l'opzione di network-level blocking via "
        "declarativeNetRequest fornisce un livello aggiuntivo "
        "di hardening contro i tunnel BitM+.")

    add_heading(doc, "6.2 Limiti del lavoro", level=2)

    add_par(doc,
        "Pur ritenendo i risultati ottenuti soddisfacenti "
        "rispetto agli obiettivi posti, è doveroso esplicitare "
        "i limiti del lavoro per non offrirne una "
        "rappresentazione idealizzata.")

    add_par(doc,
        "Il primo limite è legato alla natura della suite "
        "di test: pur coprendo le tre famiglie di stack BitM "
        "documentate in letteratura, i 49 casi sono di "
        "natura sintetica e non rappresentano un campione "
        "statisticamente significativo del traffico reale. "
        "Una valutazione su traffico produttivo, con "
        "raccolta consensuale di sessioni reali e "
        "etichettatura manuale di un dataset di centinaia "
        "o migliaia di sessioni, è il prossimo passo "
        "naturale del lavoro.")

    add_par(doc,
        "Il secondo limite riguarda la dipendenza dei "
        "marker BitM/BitM+ dalla qualità del fingerprint "
        "raccolto dal client. Un attaccante sufficientemente "
        "sofisticato può rimuovere selettivamente i campi "
        "del payload che rivelano la sua infrastruttura, "
        "depotenziando di conseguenza il rilevamento. La "
        "difesa rimane efficace contro toolkit out-of-the-box "
        "(Evilginx, Modlishka, EvilProxy nelle loro "
        "configurazioni standard), ma è meno robusta "
        "contro avversari capaci di customizzazione "
        "approfondita.")

    add_par(doc,
        "Il terzo limite è di natura architetturale: il "
        "broadcaster real-time è in-process e quindi "
        "single-worker, una limitazione che restringe la "
        "scalabilità orizzontale del servizio. La promozione "
        "del trasporto a Redis pub/sub è documentata come "
        "evoluzione necessaria per deployment multi-worker, "
        "ma non è stata implementata in questo lavoro.")

    add_par(doc,
        "Il quarto limite è economico: l'utilizzo del "
        "backend Anthropic comporta un costo per richiesta "
        "che, per quanto contenuto (dell'ordine del decimo "
        "di centesimo per richiesta con Claude Haiku 4.5), "
        "diventa significativo su volumi di traffico elevati. "
        "L'alternativa Ollama elimina il costo ricorrente "
        "ma richiede risorse di calcolo locali (idealmente "
        "una GPU) che non tutti i deployment possono "
        "permettersi.")

    add_heading(doc, "6.3 Direzioni di ricerca future", level=2)

    add_par(doc,
        "Il lavoro presentato si presta a numerose "
        "estensioni, alcune delle quali sono già "
        "infrastrutturate nel repository ma non ancora "
        "oggetto di valutazione su larga scala.")

    add_par(doc,
        "La direzione probabilmente più promettente è il "
        "fine-tuning specializzato del modello Llama 3.1 7B "
        "tramite tecniche LoRA (Low-Rank Adaptation), "
        "utilizzando il dataset estratto dal log degli "
        "eventi (bitm_events.jsonl). L'infrastruttura per "
        "convertire il log in dataset ChatML "
        "(training/build_dataset.py) e per eseguire il "
        "training (training/train_lora.py) è già presente "
        "nel repository. L'ipotesi di lavoro è che un "
        "modello specializzato sul task di rilevamento "
        "BitM possa raggiungere o superare le prestazioni "
        "di Claude Haiku con costi di inferenza inferiori "
        "di un ordine di grandezza, in modo simile a quanto "
        "documentato in letteratura per altri task verticali "
        "di sicurezza.")

    add_par(doc,
        "Una seconda direzione è l'integrazione di feed di "
        "threat intelligence esterni — passive DNS per "
        "domini freschi, certificate transparency logs per "
        "certificati TLS appena emessi, blocklist Tor "
        "aggiornate dinamicamente — che colmerebbero il gap "
        "di rilevamento contro avversari capaci di "
        "neutralizzare i marker BitM. La struttura del "
        "campo ip_meta del payload è già preparata per "
        "ricevere hint da feed esterni.")

    add_par(doc,
        "Una terza direzione è il porting dell'analisi di "
        "traiettoria a un modello di sequence labeling "
        "(LSTM o Transformer dedicato) addestrato su "
        "sessioni etichettate. Il pattern attualmente "
        "riconosciuto dal layer Trajectory — "
        "panic_password_change, direct_admin_access, "
        "rapid_navigation — potrebbe essere ampliato a "
        "decine di pattern via training supervisionato.")

    add_par(doc,
        "Una quarta direzione, di natura più applicativa, "
        "è la realizzazione di plugin nativi per i "
        "principali framework web (WordPress, Joomla, "
        "Magento, Shopify) che integrino il backend "
        "BitM-LLM in modalità one-click. La struttura "
        "REST dell'API, già minimalista, si presta "
        "naturalmente a questo tipo di integrazione.")

    add_par(doc,
        "Infine, sul versante metodologico, sarebbe "
        "interessante condurre uno studio comparativo con i "
        "principali sistemi commerciali di bot management "
        "(Cloudflare Turnstile, Akamai Bot Manager, "
        "DataDome) su un dataset di attacchi BitM+ "
        "reali, per quantificare il vantaggio offerto "
        "dall'approccio integrato proposto rispetto alle "
        "soluzioni di mercato basate prevalentemente su "
        "machine learning statico.")

    add_par(doc,
        "Il problema della difesa contro il "
        "Browser-in-the-Middle è ben lungi dall'essere "
        "risolto, e la rapidità con cui i toolkit di "
        "attacco evolvono — sostituendo le proprie firme, "
        "diversificando i propri stack, integrando bypass "
        "di nuovi meccanismi MFA — impone una postura "
        "difensiva altrettanto adattiva. L'architettura "
        "presentata in questo lavoro, integrando "
        "deterministico e probabilistico in un'unica "
        "pipeline e demandando le valutazioni qualitative a "
        "un Large Language Model, costituisce un passo in "
        "questa direzione, e si offre alla comunità "
        "scientifica e operativa come base di partenza per "
        "ulteriori miglioramenti.")

    add_page_break(doc)


def bibliografia(doc):
    add_heading(doc, "Bibliografia", level=1)

    refs = [
        '[1] F. Tommasi, C. Catalano, I. Taurino, "Browser-in-the-Middle (BitM) attack", '
        'International Journal of Information Security, vol. 21, pp. 179–189, 2022. '
        'DOI: 10.1007/s10207-021-00548-5.',

        '[2] C. Catalano, A. Pagano, A. Piccinno, A. Stamerra, "Cartoons to Improve '
        'Cyber Security Education: Snow White in Browser in the Middle", '
        'Proceedings of IS-EUD 2023, 9th International Symposium on End-User Development, '
        'Cagliari, Italy, June 2023.',

        '[3] D. Tzschoppe et al., "Apache Guacamole as a Browser-in-the-Middle Attack '
        'Vector: Analysis and Countermeasures", Journal of Computer Virology and '
        'Hacking Techniques, 2023.',

        '[4] V. Jagannath, "Browser-in-the-Middle Attacks: A Comprehensive Analysis '
        'and Countermeasures", Security and Privacy, Wiley, 2024.',

        '[5] C. Catalano et al., "BitM+: Bypassing WebAuthn through Browser-in-the-Middle '
        'with reflected XSS payloads", Journal of Computer Virology and Hacking '
        'Techniques, 2025.',

        '[6] FIDO Alliance, "Web Authentication: An API for accessing Public Key '
        'Credentials, Level 2", W3C Recommendation, 2021. '
        'URL: https://www.w3.org/TR/webauthn-2/.',

        '[7] OWASP Foundation, "OWASP Top Ten Web Application Security Risks", 2021. '
        'URL: https://owasp.org/www-project-top-ten/.',

        '[8] Anthropic, "Claude Models Documentation", 2024. '
        'URL: https://docs.anthropic.com/.',

        '[9] Meta AI, "LLaMA 3 Technical Report", 2024.',

        '[10] E. J. Hu et al., "LoRA: Low-Rank Adaptation of Large Language Models", '
        'International Conference on Learning Representations (ICLR), 2022.',

        '[11] J. Dai et al., "Phishing detection using a combination of URL features '
        'and content-based detection", Computers & Security, vol. 67, 2017.',

        '[12] S. Ramirez, "FastAPI: Modern, fast (high-performance), web framework for '
        'building APIs with Python", 2018. URL: https://fastapi.tiangolo.com/.',

        '[13] noVNC project, "HTML VNC client". URL: https://github.com/novnc/noVNC.',

        '[14] Apache Software Foundation, "Apache Guacamole: Clientless remote desktop '
        'gateway". URL: https://guacamole.apache.org/.',

        '[15] R. Petersen, D. Santos, M. C. Smith, K. A. Wetzel, G. Witte, "Workforce '
        'Framework for Cybersecurity (NICE Framework)", NIST Special Publication '
        '800-181, 2020.',

        '[16] ENISA, "European Cybersecurity Skills Framework (ECSF) — User Manual", '
        '2022. URL: https://www.enisa.europa.eu/.',

        '[17] A. Rashid et al., "Scoping the Cyber Security Body of Knowledge", '
        'IEEE Security & Privacy, vol. 16, no. 3, pp. 96–102, 2018.',

        '[18] Redis Ltd., "Redis Documentation". URL: https://redis.io/docs/.',

        '[19] MaxMind Inc., "GeoLite2 Free Geolocation Data". '
        'URL: https://dev.maxmind.com/geoip/geolite2-free-geolocation-data.',

        '[20] Google LLC, "Chrome Extensions Manifest V3 Documentation". '
        'URL: https://developer.chrome.com/docs/extensions/develop/migrate.',
    ]

    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(r)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    add_page_break(doc)


def ringraziamenti(doc):
    add_heading(doc, "Ringraziamenti", level=1)

    add_par(doc,
        "Al termine di questo percorso di studi, sento il "
        "bisogno di esprimere la mia riconoscenza a tutte le "
        "persone che, in modo diretto o indiretto, hanno "
        "contribuito al raggiungimento di questo traguardo.")

    add_par(doc,
        "Desidero ringraziare in primo luogo il mio relatore, "
        "il Chiar.mo Prof. [Nome del Relatore], per la "
        "disponibilità, la pazienza e la competenza con cui "
        "ha guidato lo sviluppo di questa tesi. I suoi "
        "suggerimenti, sempre puntuali e stimolanti, hanno "
        "contribuito in modo determinante a indirizzare il "
        "lavoro nella sua forma attuale.")

    add_par(doc,
        "Un sentito ringraziamento va al Dipartimento di "
        "Informatica dell'Università degli Studi di Bari "
        "Aldo Moro, ai docenti e al personale tutto, per "
        "aver reso possibile un percorso di studi che si è "
        "rivelato non solo formativo sotto il profilo "
        "tecnico, ma anche di crescita personale e umana.")

    add_par(doc,
        "Ringrazio i miei genitori per il sostegno costante "
        "e silenzioso che mi hanno offerto in tutti questi "
        "anni, per aver creduto in me anche nei momenti in "
        "cui io stesso vacillavo, per essere stati la "
        "presenza solida su cui ho potuto sempre contare. "
        "Senza il loro supporto, nulla di tutto questo "
        "sarebbe stato possibile.")

    add_par(doc,
        "Ai miei amici e compagni di corso, con cui ho "
        "condiviso fatiche, soddisfazioni, sessioni "
        "interminabili di studio e momenti di leggerezza, "
        "va un ringraziamento sincero per aver reso questi "
        "anni di università un'esperienza che ricorderò con "
        "affetto.")

    add_par(doc,
        "Infine, un pensiero particolare va a chi, pur non "
        "essendo presente fisicamente in questa giornata, "
        "ha lasciato in me un'impronta che mi accompagnerà "
        "nel percorso futuro. Questo lavoro è anche per voi.")

    add_par(doc, "")
    add_par(doc, "Bari, [data della discussione]",
            italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=False)
    add_par(doc, "Gabriele",
            italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
            first_line_indent=False)


# ============================================================================
# MAIN
# ============================================================================

def main():
    doc = init_doc_from_template()
    setup_header_footer(doc)

    frontespizio(doc)
    dedica(doc)
    indice(doc)
    introduzione(doc)
    capitolo1_stato_arte(doc)
    capitolo2_problema(doc)
    capitolo3_progettazione(doc)
    capitolo4_implementazione(doc)
    capitolo5_sperimentazione(doc)
    capitolo6_conclusioni(doc)
    bibliografia(doc)
    ringraziamenti(doc)

    out = Path(__file__).parent / "tesi_BitM_LLM.docx"
    doc.save(str(out))
    print(f"Documento generato a partire da Template_Tesi.docx: {out}")


if __name__ == "__main__":
    main()
